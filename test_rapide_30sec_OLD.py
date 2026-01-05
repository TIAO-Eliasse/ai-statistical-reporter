"""
TEST ULTRA-RAPIDE - 30 SECONDES
VÃ©rifie que le module table_formatter fonctionne avec VOS donnÃ©es
"""

import pandas as pd
from table_formatter import TableFormatter
from docx import Document
import sys

print("\n" + "="*70)
print("ğŸ§ª TEST RAPIDE - TABLE FORMATTER")
print("="*70 + "\n")

# ========== Ã‰TAPE 1 : VOS DONNÃ‰ES ==========
print("ğŸ“Š Ã‰tape 1/4 : Chargement des donnÃ©es...")

df = pd.DataFrame({
    'nom': ['Alice', 'Bob', 'Charlie', 'Diana', 'Eve', 'Frank'],
    'age': [25, 30, 35, 28, 32, 29],
    'salaire': [45000, 52000, 48000, 51000, 55000, 47000],
    'ville': ['Paris', 'Lyon', 'Marseille', 'Paris', 'Lyon', 'Toulouse']
})

print(f"âœ… DonnÃ©es chargÃ©es : {len(df)} lignes, {len(df.columns)} colonnes")


# ========== Ã‰TAPE 2 : TEST HTML ==========
print("\nğŸ“„ Ã‰tape 2/4 : Test export HTML...")

try:
    formatter = TableFormatter(style='professional')
    
    # Test 1 : Statistiques descriptives
    stats_html = formatter.dataframe_to_html(
        df.describe().T,
        title="Statistiques descriptives"
    )
    
    # Test 2 : Moyennes par ville
    ville_html = formatter.dataframe_to_html(
        df.groupby('ville').mean().round(2),
        title="Moyennes par ville",
        include_index=True
    )
    
    print("âœ… HTML gÃ©nÃ©rÃ© avec succÃ¨s")
    print(f"   - Taille stats : {len(stats_html)} caractÃ¨res")
    print(f"   - Taille ville : {len(ville_html)} caractÃ¨res")
    
except Exception as e:
    print(f"âŒ ERREUR HTML : {e}")
    sys.exit(1)


# ========== Ã‰TAPE 3 : TEST WORD ==========
print("\nğŸ“ Ã‰tape 3/4 : Test export Word...")

try:
    doc = Document()
    doc.add_heading('Test Rapport - Table Formatter', 0)
    doc.add_paragraph('Ce document dÃ©montre les tableaux professionnels.')
    
    # Ajouter tableaux
    doc = formatter.dataframe_to_docx_table(
        doc,
        df.describe().T,
        title="âœ… Statistiques descriptives (nouveau format)"
    )
    
    doc = formatter.dataframe_to_docx_table(
        doc,
        df.groupby('ville').mean().round(2),
        title="âœ… Moyennes par ville (nouveau format)",
        include_index=True
    )
    
    # Sauvegarder
    doc.save('TEST_TABLEAUX_PROFESSIONNEL.docx')
    
    print("âœ… Document Word crÃ©Ã©")
    print("   ğŸ“ Fichier : TEST_TABLEAUX_PROFESSIONNEL.docx")
    
except Exception as e:
    print(f"âŒ ERREUR Word : {e}")
    sys.exit(1)


# ========== Ã‰TAPE 4 : COMPARAISON AVANT/APRÃˆS ==========
print("\nğŸ” Ã‰tape 4/4 : Comparaison AVANT/APRÃˆS...")

print("\n" + "-"*70)
print("âŒ AVANT (texte brut) :")
print("-"*70)
print(str(df.describe())[:200] + "...")

print("\n" + "-"*70)
print("âœ… APRÃˆS (tableau HTML) :")
print("-"*70)
print(stats_html[:300] + "...")


# ========== RÃ‰SUMÃ‰ ==========
print("\n" + "="*70)
print("âœ… TOUS LES TESTS PASSÃ‰S !")
print("="*70)
print()
print("ğŸ“‹ PROCHAINES ACTIONS :")
print()
print("1. âœ… Ouvrez TEST_TABLEAUX_PROFESSIONNEL.docx")
print("   â†’ Les tableaux doivent Ãªtre magnifiques (couleurs, bordures)")
print()
print("2. âœ… Comparez avec vos anciens rapports")
print("   â†’ Avant : texte brut illisible")
print("   â†’ AprÃ¨s : tableaux professionnels")
print()
print("3. âœ… Si tout est OK, passez Ã  l'intÃ©gration dans votre code")
print("   â†’ Consultez EXEMPLE_CORRECTIONS.md")
print()
print("="*70)
print()

print("ğŸ‰ TEST RÃ‰USSI ! Votre module fonctionne parfaitement.")
print()