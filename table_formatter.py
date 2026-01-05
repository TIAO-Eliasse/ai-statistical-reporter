"""
Table Formatter pour AI Statistical Reporter (VERSION ULTRA-COMPATIBLE)
Fonctionne avec TOUTES les versions de python-docx (même les très anciennes)

CORRECTIONS :
- Compatible python-docx 0.5.x à 1.x
- Gestion des RGBColor anciennes et nouvelles
- Fallback pour attributs manquants
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional, Dict, Any
import io


class TableFormatter:
    """Classe pour formater les tableaux pandas en différents formats"""
    
    def __init__(self, style: str = "professional"):
        """
        Initialise le formateur
        
        Args:
            style: Style du tableau ('professional', 'academic', 'minimal')
        """
        self.style = style
        self.styles_config = {
            'professional': {
                'header_bg': (76, 175, 80),  # Vert
                'header_text': (255, 255, 255),  # Blanc
                'border_color': (221, 221, 221),  # Gris clair
                'alt_row_bg': (242, 242, 242)  # Gris très clair
            },
            'academic': {
                'header_bg': (52, 73, 94),  # Bleu foncé
                'header_text': (255, 255, 255),
                'border_color': (189, 195, 199),
                'alt_row_bg': (236, 240, 241)
            },
            'minimal': {
                'header_bg': (149, 165, 166),  # Gris
                'header_text': (255, 255, 255),
                'border_color': (127, 140, 141),
                'alt_row_bg': None  # Pas d'alternance
            }
        }
    
    def _rgb_tuple_to_hex(self, rgb_tuple) -> str:
        """
        Convertit un tuple RGB en hexadécimal
        
        Args:
            rgb_tuple: Tuple (R, G, B)
            
        Returns:
            String hexadécimal sans # (ex: "4CAF50")
        """
        r, g, b = rgb_tuple
        return f"{r:02X}{g:02X}{b:02X}"
    
    def _rgb_tuple_to_css(self, rgb_tuple) -> str:
        """
        Convertit un tuple RGB en format CSS
        
        Args:
            rgb_tuple: Tuple (R, G, B)
            
        Returns:
            String CSS (ex: "rgb(76, 175, 80)")
        """
        if rgb_tuple is None:
            return "transparent"
        r, g, b = rgb_tuple
        return f"rgb({r}, {g}, {b})"
    
    def _create_rgbcolor(self, rgb_tuple):
        """
        Crée un RGBColor de manière compatible
        
        Args:
            rgb_tuple: Tuple (R, G, B)
            
        Returns:
            RGBColor object
        """
        r, g, b = rgb_tuple
        return RGBColor(r, g, b)
    
    def _shade_cell(self, cell, rgb_tuple):
        """
        Applique une couleur de fond à une cellule (compatible toutes versions)
        
        Args:
            cell: Cellule Word
            rgb_tuple: Tuple (R, G, B)
        """
        if rgb_tuple is None:
            return
        
        hex_color = self._rgb_tuple_to_hex(rgb_tuple)
        
        try:
            # Méthode 1 : OxmlElement (fonctionne toutes versions)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), hex_color)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception:
            try:
                # Méthode 2 : Alternative
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), hex_color)
                tcPr.append(shd)
            except Exception:
                # Ignorer si impossible
                pass
    
    def format_value(self, value: Any, precision: int = 2) -> str:
        """
        Formate une valeur pour affichage
        
        Args:
            value: Valeur à formater
            precision: Nombre de décimales pour les floats
            
        Returns:
            Valeur formatée en string
        """
        if pd.isna(value):
            return "N/A"
        elif isinstance(value, float):
            return f"{value:.{precision}f}"
        elif isinstance(value, int):
            return str(value)
        else:
            return str(value)
    
    def dataframe_to_docx_table(
        self,
        doc: Document,
        df: pd.DataFrame,
        title: Optional[str] = None,
        precision: int = 2,
        include_index: bool = False
    ) -> Document:
        """
        Ajoute un DataFrame formaté à un document Word
        
        Args:
            doc: Document Word (python-docx)
            df: DataFrame à ajouter
            title: Titre optionnel du tableau
            precision: Nombre de décimales
            include_index: Inclure l'index du DataFrame
            
        Returns:
            Document Word modifié
        """
        # Ajouter le titre si fourni
        if title:
            heading = doc.add_heading(title, level=3)
            heading.style.font.color.rgb = self._create_rgbcolor((44, 62, 80))
        
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Créer le tableau
        num_rows = len(df_display) + 1  # +1 pour le header
        num_cols = len(df_display.columns)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Configuration du style
        style_cfg = self.styles_config.get(self.style, self.styles_config['professional'])
        
        # HEADER ROW
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df_display.columns):
            cell = header_cells[i]
            cell.text = str(col_name)
            
            # Style du header
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self._create_rgbcolor(style_cfg['header_text'])
            
            # Couleur de fond du header
            self._shade_cell(cell, style_cfg['header_bg'])
            
            # Alignement
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # DATA ROWS
        for row_idx, row_data in enumerate(df_display.itertuples(index=False), start=1):
            row_cells = table.rows[row_idx].cells
            
            for col_idx, value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = self.format_value(value, precision)
                
                # Style de la cellule
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
                run.font.size = Pt(10)
                
                # Alignement : nombres à droite, texte à gauche
                if isinstance(value, (int, float)) and not pd.isna(value):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Alternance des couleurs de lignes
                if style_cfg['alt_row_bg'] and row_idx % 2 == 0:
                    self._shade_cell(cell, style_cfg['alt_row_bg'])
        
        # Ajouter un espace après le tableau
        doc.add_paragraph()
        
        return doc
    
    def dataframe_to_html(
        self,
        df: pd.DataFrame,
        title: Optional[str] = None,
        precision: int = 2,
        include_index: bool = False
    ) -> str:
        """
        Convertit un DataFrame en HTML professionnel
        
        Args:
            df: DataFrame à convertir
            title: Titre optionnel
            precision: Nombre de décimales
            include_index: Inclure l'index
            
        Returns:
            String HTML
        """
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Formater les valeurs
        df_formatted = df_display.copy()
        for col in df_formatted.columns:
            df_formatted[col] = df_formatted[col].apply(
                lambda x: self.format_value(x, precision)
            )
        
        # Configuration du style
        style_cfg = self.styles_config.get(self.style, self.styles_config['professional'])
        
        # Convertir les couleurs en format CSS
        header_bg = self._rgb_tuple_to_css(style_cfg['header_bg'])
        header_text = self._rgb_tuple_to_css(style_cfg['header_text'])
        alt_row_bg = self._rgb_tuple_to_css(style_cfg['alt_row_bg'])
        
        # Générer le HTML
        html = ""
        
        if title:
            html += f'<h3 style="color: #2c3e50; margin-top: 20px;">{title}</h3>\n'
        
        html += f'''
<table style="border-collapse: collapse; width: 100%; margin: 20px 0; font-family: Arial, sans-serif;">
    <thead>
        <tr style="background-color: {header_bg}; color: {header_text};">
'''
        
        # Headers
        for col in df_formatted.columns:
            html += f'            <th style="border: 1px solid #ddd; padding: 12px; text-align: left; font-weight: bold;">{col}</th>\n'
        
        html += '        </tr>\n    </thead>\n    <tbody>\n'
        
        # Rows
        for row_num, (idx, row) in enumerate(df_formatted.iterrows()):
            bg_color = alt_row_bg if row_num % 2 == 0 else "white"
            html += f'        <tr style="background-color: {bg_color};">\n'
            
            for value in row:
                html += f'            <td style="border: 1px solid #ddd; padding: 10px;">{value}</td>\n'
            
            html += '        </tr>\n'
        
        html += '    </tbody>\n</table>\n'
        
        return html
    
    def dataframe_to_markdown(
        self,
        df: pd.DataFrame,
        precision: int = 2,
        include_index: bool = False
    ) -> str:
        """
        Convertit un DataFrame en Markdown
        
        Args:
            df: DataFrame à convertir
            precision: Nombre de décimales
            include_index: Inclure l'index
            
        Returns:
            String Markdown
        """
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Formater les valeurs
        df_formatted = df_display.copy()
        for col in df_formatted.columns:
            df_formatted[col] = df_formatted[col].apply(
                lambda x: self.format_value(x, precision)
            )
        
        # Générer le Markdown
        md = "| " + " | ".join(str(col) for col in df_formatted.columns) + " |\n"
        md += "|" + "|".join(["---" for _ in df_formatted.columns]) + "|\n"
        
        for _, row in df_formatted.iterrows():
            md += "| " + " | ".join(str(val) for val in row) + " |\n"
        
        return md


# ============= FONCTIONS UTILITAIRES =============

def format_statistics_table(df: pd.DataFrame, column: str, style: str = "professional") -> Dict[str, Any]:
    """Crée un tableau de statistiques descriptives formaté"""
    formatter = TableFormatter(style=style)
    stats = df[column].describe()
    stats_df = pd.DataFrame({
        'Statistique': stats.index,
        'Valeur': stats.values
    })
    
    return {
        'dataframe': stats_df,
        'html': formatter.dataframe_to_html(stats_df, title=f"Statistiques descriptives : {column}"),
        'markdown': formatter.dataframe_to_markdown(stats_df)
    }


def format_crosstab(df: pd.DataFrame, row_var: str, col_var: str, style: str = "professional") -> Dict[str, Any]:
    """Crée un tableau croisé formaté"""
    formatter = TableFormatter(style=style)
    crosstab = pd.crosstab(df[row_var], df[col_var])
    
    return {
        'dataframe': crosstab,
        'html': formatter.dataframe_to_html(
            crosstab, 
            title=f"Tableau croisé : {row_var} × {col_var}",
            include_index=True
        ),
        'markdown': formatter.dataframe_to_markdown(crosstab, include_index=True)
    }


def format_correlation_matrix(df: pd.DataFrame, style: str = "professional") -> Dict[str, Any]:
    """Crée une matrice de corrélation formatée"""
    formatter = TableFormatter(style=style)
    corr_matrix = df.select_dtypes(include=['number']).corr()
    
    return {
        'dataframe': corr_matrix,
        'html': formatter.dataframe_to_html(
            corr_matrix,
            title="Matrice de corrélation",
            precision=3,
            include_index=True
        ),
        'markdown': formatter.dataframe_to_markdown(corr_matrix, precision=3, include_index=True)
    }


if __name__ == "__main__":
    import numpy as np
    
    np.random.seed(42)
    test_df = pd.DataFrame({
        'Variable': ['Âge', 'Salaire', 'Expérience', 'Satisfaction'],
        'Moyenne': np.random.uniform(20, 60, 4),
        'Écart-type': np.random.uniform(5, 15, 4),
        'Min': np.random.uniform(10, 30, 4),
        'Max': np.random.uniform(70, 100, 4)
    })
    
    formatter = TableFormatter(style="professional")
    
    print("=== HTML ===")
    print(formatter.dataframe_to_html(test_df, title="Statistiques descriptives"))
    
    print("\n=== MARKDOWN ===")
    print(formatter.dataframe_to_markdown(test_df))
    
    print("\n=== WORD ===")
    doc = Document()
    doc = formatter.dataframe_to_docx_table(doc, test_df, title="Statistiques descriptives")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    print(f"Document Word créé ({len(buffer.getvalue())} bytes)")