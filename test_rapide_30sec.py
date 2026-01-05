"""
TEST ULTRA-RAPIDE - 30 SECONDES (VERSION CORRIGÉE)
Vérifie que le module table_formatter fonctionne avec VOS données
"""

import pandas as pd
from table_formatter import TableFormatter
from docx import Document
import sys

print("\n" + "="*70)
print("🧪 TEST RAPIDE - TABLE FORMATTER")
print("="*70 + "\n")

# ========== ÉTAPE 1 : VOS DONNÉES ==========
print("📊 Étape 1/4 : Chargement des données...")

df = pd.DataFrame({
    'nom': ['Alice', 'Bob', 'Charlie', 'Diana', 'Eve', 'Frank'],
    'age': [25, 30, 35, 28, 32, 29],
    'salaire': [45000, 52000, 48000, 51000, 55000, 47000],
    'ville': ['Paris', 'Lyon', 'Marseille', 'Paris', 'Lyon', 'Toulouse']
})

print(f"✅ Données chargées : {len(df)} lignes, {len(df.columns)} colonnes")


# ========== ÉTAPE 2 : TEST HTML ==========
print("\n📄 Étape 2/4 : Test export HTML...")

try:
    formatter = TableFormatter(style='professional')
    
    # Test 1 : Statistiques descriptives
    stats_df = df.describe().T
    stats_html = formatter.dataframe_to_html(
        stats_df,
        title="Statistiques descriptives"
    )
    
    # Test 2 : Moyennes par ville (colonnes numériques seulement)
    ville_df = df.groupby('ville')[['age', 'salaire']].mean().round(2)
    ville_html = formatter.dataframe_to_html(
        ville_df,
        title="Moyennes par ville",
        include_index=True
    )
    
    print("✅ HTML généré avec succès")
    print(f"   - Taille stats : {len(stats_html)} caractères")
    print(f"   - Taille ville : {len(ville_html)} caractères")
    
except Exception as e:
    print(f"❌ ERREUR HTML : {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)


# ========== ÉTAPE 3 : TEST WORD ==========
print("\n📝 Étape 3/4 : Test export Word...")

try:
    doc = Document()
    doc.add_heading('Test Rapport - Table Formatter', 0)
    doc.add_paragraph('Ce document démontre les tableaux professionnels.')
    
    # Ajouter tableaux
    stats_df = df.describe().T
    doc = formatter.dataframe_to_docx_table(
        doc,
        stats_df,
        title="✅ Statistiques descriptives (nouveau format)"
    )
    
    ville_df = df.groupby('ville')[['age', 'salaire']].mean().round(2)
    doc = formatter.dataframe_to_docx_table(
        doc,
        ville_df,
        title="✅ Moyennes par ville (nouveau format)",
        include_index=True
    )
    
    # Sauvegarder
    doc.save('TEST_TABLEAUX_PROFESSIONNEL.docx')
    
    print("✅ Document Word créé")
    print("   📁 Fichier : TEST_TABLEAUX_PROFESSIONNEL.docx")
    
except Exception as e:
    print(f"❌ ERREUR Word : {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)


# ========== ÉTAPE 4 : COMPARAISON AVANT/APRÈS ==========
print("\n🔍 Étape 4/4 : Comparaison AVANT/APRÈS...")

print("\n" + "-"*70)
print("❌ AVANT (texte brut) :")
print("-"*70)
print(str(df.describe())[:200] + "...")

print("\n" + "-"*70)
print("✅ APRÈS (tableau HTML) :")
print("-"*70)
print(stats_html[:300] + "...")


# ========== RÉSUMÉ ==========
print("\n" + "="*70)
print("✅ TOUS LES TESTS PASSÉS !")
print("="*70)
print()
print("📋 PROCHAINES ACTIONS :")
print()
print("1. ✅ Ouvrez TEST_TABLEAUX_PROFESSIONNEL.docx")
print("   → Les tableaux doivent être magnifiques (couleurs, bordures)")
print()
print("2. ✅ Comparez avec vos anciens rapports")
print("   → Avant : texte brut illisible")
print("   → Après : tableaux professionnels")
print()
print("3. ✅ Si tout est OK, passez à l'intégration dans votre code")
print("   → Consultez EXEMPLE_CORRECTIONS.md")
print()
print("="*70)
print()

print("🎉 TEST RÉUSSI ! Votre module fonctionne parfaitement.")
print()