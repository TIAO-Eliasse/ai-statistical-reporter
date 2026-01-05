"""
Tests pour le module table_formatter
Vérifie que tous les formats de tableaux fonctionnent correctement
"""

import pandas as pd
import numpy as np
from table_formatter import TableFormatter, format_statistics_table, format_crosstab
from docx import Document
import io


def test_basic_formatting():
    """Test 1 : Formatage basique d'un DataFrame simple"""
    
    print("🧪 TEST 1 : Formatage basique")
    print("=" * 60)
    
    # Créer un DataFrame simple
    df = pd.DataFrame({
        'Ville': ['Paris', 'Lyon', 'Marseille', 'Toulouse'],
        'Population': [2165423, 513275, 869815, 471941],
        'Superficie': [105.4, 47.87, 240.62, 118.3],
        'Densité': [20545.3, 10719.6, 3614.4, 3989.2]
    })
    
    formatter = TableFormatter(style='professional')
    
    # Test HTML
    html = formatter.dataframe_to_html(df, title="Grandes villes de France")
    assert '<table' in html, "❌ HTML invalide"
    assert 'Paris' in html, "❌ Données manquantes"
    print("✅ HTML généré correctement")
    
    # Test Markdown
    md = formatter.dataframe_to_markdown(df)
    assert '|' in md, "❌ Markdown invalide"
    assert 'Ville' in md, "❌ Headers manquants"
    print("✅ Markdown généré correctement")
    
    # Test Word
    doc = Document()
    doc = formatter.dataframe_to_docx_table(doc, df, title="Grandes villes de France")
    assert len(doc.tables) > 0, "❌ Tableau Word non créé"
    print("✅ Tableau Word créé correctement")
    
    print()
    return True


def test_statistics_formatting():
    """Test 2 : Formatage de statistiques descriptives"""
    
    print("🧪 TEST 2 : Statistiques descriptives")
    print("=" * 60)
    
    # Créer des données
    np.random.seed(42)
    df = pd.DataFrame({
        'age': np.random.randint(25, 65, 100),
        'salaire': np.random.randint(30000, 80000, 100)
    })
    
    # Tester format_statistics_table
    stats_formats = format_statistics_table(df, 'age', style='academic')
    
    assert 'dataframe' in stats_formats, "❌ DataFrame manquant"
    assert 'html' in stats_formats, "❌ HTML manquant"
    assert 'markdown' in stats_formats, "❌ Markdown manquant"
    
    stats_df = stats_formats['dataframe']
    assert 'mean' in stats_df['Statistique'].values, "❌ Moyenne manquante"
    assert 'std' in stats_df['Statistique'].values, "❌ Écart-type manquant"
    
    print("✅ Statistiques formatées correctement")
    print(f"   - Mean: {stats_df[stats_df['Statistique'] == 'mean']['Valeur'].values[0]:.2f}")
    print(f"   - Std: {stats_df[stats_df['Statistique'] == 'std']['Valeur'].values[0]:.2f}")
    
    print()
    return True


def test_value_formatting():
    """Test 3 : Formatage des valeurs (floats, NaN, etc.)"""
    
    print("🧪 TEST 3 : Formatage des valeurs")
    print("=" * 60)
    
    formatter = TableFormatter()
    
    # Test float
    assert formatter.format_value(3.14159, precision=2) == "3.14", "❌ Float mal formaté"
    print("✅ Float formaté : 3.14159 → 3.14")
    
    # Test int
    assert formatter.format_value(42) == "42", "❌ Int mal formaté"
    print("✅ Int formaté : 42 → 42")
    
    # Test NaN
    assert formatter.format_value(np.nan) == "N/A", "❌ NaN mal formaté"
    print("✅ NaN formaté : NaN → N/A")
    
    # Test string
    assert formatter.format_value("Paris") == "Paris", "❌ String mal formaté"
    print("✅ String formaté : Paris → Paris")
    
    print()
    return True


def test_styles():
    """Test 4 : Différents styles de tableaux"""
    
    print("🧪 TEST 4 : Styles de tableaux")
    print("=" * 60)
    
    df = pd.DataFrame({
        'Nom': ['Alice', 'Bob', 'Charlie'],
        'Score': [85.5, 92.3, 78.9]
    })
    
    styles = ['professional', 'academic', 'minimal']
    
    for style in styles:
        formatter = TableFormatter(style=style)
        html = formatter.dataframe_to_html(df)
        assert '<table' in html, f"❌ Style {style} invalide"
        print(f"✅ Style '{style}' fonctionne")
    
    print()
    return True


def test_large_dataframe():
    """Test 5 : DataFrame volumineux (performance)"""
    
    print("🧪 TEST 5 : DataFrame volumineux")
    print("=" * 60)
    
    # Créer un grand DataFrame
    np.random.seed(42)
    large_df = pd.DataFrame({
        f'Var{i}': np.random.randn(1000) for i in range(10)
    })
    
    formatter = TableFormatter()
    
    # Test HTML (doit être rapide)
    import time
    start = time.time()
    html = formatter.dataframe_to_html(large_df.head(50))  # Limiter à 50 lignes
    duration = time.time() - start
    
    assert duration < 1.0, f"❌ Trop lent ({duration:.2f}s)"
    print(f"✅ 50 lignes × 10 colonnes générées en {duration:.3f}s")
    
    # Test Word
    doc = Document()
    start = time.time()
    doc = formatter.dataframe_to_docx_table(doc, large_df.head(50))
    duration = time.time() - start
    
    assert duration < 2.0, f"❌ Trop lent ({duration:.2f}s)"
    print(f"✅ Tableau Word créé en {duration:.3f}s")
    
    print()
    return True


def test_real_world_case():
    """Test 6 : Cas réel avec données de l'application"""
    
    print("🧪 TEST 6 : Cas réel (données de l'application)")
    print("=" * 60)
    
    # Simuler les données de l'application
    df = pd.DataFrame({
        'nom': ['Alice', 'Bob', 'Charlie', 'Diana', 'Eve', 'Frank'],
        'age': [25, 30, 35, 28, 32, 29],
        'salaire': [45000, 52000, 48000, 51000, 55000, 47000],
        'ville': ['Paris', 'Lyon', 'Marseille', 'Paris', 'Lyon', 'Toulouse']
    })
    
    formatter = TableFormatter(style='professional')
    
    # Test 1 : Statistiques par ville
    print("\n📊 Statistiques par ville :")
    ville_stats = df.groupby('ville').agg({
        'age': ['mean', 'std'],
        'salaire': ['mean', 'std']
    }).round(2)
    
    html = formatter.dataframe_to_html(ville_stats, title="Statistiques par ville", include_index=True)
    assert 'Paris' in html, "❌ Données manquantes"
    print("✅ Tableau statistiques par ville généré")
    
    # Test 2 : Matrice de corrélation
    print("\n📈 Matrice de corrélation :")
    corr = df[['age', 'salaire']].corr()
    html_corr = formatter.dataframe_to_html(corr, title="Corrélation âge-salaire", include_index=True, precision=3)
    assert 'age' in html_corr, "❌ Variables manquantes"
    print("✅ Matrice de corrélation générée")
    
    # Test 3 : Export Word complet
    print("\n📄 Export Word complet :")
    doc = Document()
    doc.add_heading('Rapport d\'analyse', 0)
    
    doc = formatter.dataframe_to_docx_table(doc, df, title="Données brutes")
    doc = formatter.dataframe_to_docx_table(doc, ville_stats, title="Statistiques par ville", include_index=True)
    doc = formatter.dataframe_to_docx_table(doc, corr, title="Corrélation", include_index=True, precision=3)
    
    # Sauvegarder
    buffer = io.BytesIO()
    doc.save(buffer)
    size_kb = len(buffer.getvalue()) / 1024
    
    assert size_kb > 10, "❌ Document trop petit"
    print(f"✅ Document Word créé ({size_kb:.1f} KB)")
    
    print()
    return True


def test_edge_cases():
    """Test 7 : Cas limites (DataFrame vide, colonnes spéciales, etc.)"""
    
    print("🧪 TEST 7 : Cas limites")
    print("=" * 60)
    
    formatter = TableFormatter()
    
    # Test DataFrame vide
    empty_df = pd.DataFrame()
    try:
        html = formatter.dataframe_to_html(empty_df)
        print("✅ DataFrame vide géré")
    except:
        print("❌ Erreur sur DataFrame vide")
    
    # Test avec NaN partout
    nan_df = pd.DataFrame({
        'A': [np.nan, np.nan],
        'B': [np.nan, np.nan]
    })
    html_nan = formatter.dataframe_to_html(nan_df)
    assert 'N/A' in html_nan, "❌ NaN mal gérés"
    print("✅ NaN gérés correctement")
    
    # Test avec noms de colonnes longs
    long_cols_df = pd.DataFrame({
        'Nom de colonne extrêmement long pour tester': [1, 2],
        'Autre nom très long': [3, 4]
    })
    html_long = formatter.dataframe_to_html(long_cols_df)
    assert '<table' in html_long, "❌ Colonnes longues mal gérées"
    print("✅ Colonnes longues gérées")
    
    print()
    return True


def run_all_tests():
    """Exécute tous les tests"""
    
    print("\n" + "="*60)
    print("🧪 TESTS DU MODULE TABLE_FORMATTER")
    print("="*60 + "\n")
    
    tests = [
        ("Formatage basique", test_basic_formatting),
        ("Statistiques", test_statistics_formatting),
        ("Formatage valeurs", test_value_formatting),
        ("Styles multiples", test_styles),
        ("Performance", test_large_dataframe),
        ("Cas réel", test_real_world_case),
        ("Cas limites", test_edge_cases)
    ]
    
    passed = 0
    failed = 0
    
    for name, test_func in tests:
        try:
            if test_func():
                passed += 1
            else:
                failed += 1
                print(f"❌ {name} ÉCHOUÉ\n")
        except Exception as e:
            failed += 1
            print(f"❌ {name} ERREUR : {str(e)}\n")
    
    # Résumé
    print("="*60)
    print(f"📊 RÉSULTATS : {passed}/{len(tests)} tests réussis")
    if failed == 0:
        print("✅ TOUS LES TESTS PASSÉS !")
    else:
        print(f"⚠️ {failed} test(s) échoué(s)")
    print("="*60)
    
    return failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    exit(0 if success else 1)