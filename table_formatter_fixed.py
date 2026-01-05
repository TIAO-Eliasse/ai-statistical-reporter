"""
Table Formatter pour AI Statistical Reporter (VERSION CORRIGÉE)
Formate les tableaux pandas pour exports professionnels (Word, PDF, HTML)

CORRECTION : Compatible avec toutes les versions de python-docx
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional, Dict, Any
import io


class TableFormatter:
    """Classe pour formater les tableaux pandas en différents formats"""
    
    def __init__(self, style: str = "professional"):
        """
        Initialise le formateur
        
        Args:
            style: Style du tableau ('professional', 'academic', 'minimal')
        """
        self.style = style
        self.styles_config = {
            'professional': {
                'header_bg': RGBColor(76, 175, 80),  # Vert
                'header_text': RGBColor(255, 255, 255),  # Blanc
                'border_color': RGBColor(221, 221, 221),  # Gris clair
                'alt_row_bg': RGBColor(242, 242, 242)  # Gris très clair
            },
            'academic': {
                'header_bg': RGBColor(52, 73, 94),  # Bleu foncé
                'header_text': RGBColor(255, 255, 255),
                'border_color': RGBColor(189, 195, 199),
                'alt_row_bg': RGBColor(236, 240, 241)
            },
            'minimal': {
                'header_bg': RGBColor(149, 165, 166),  # Gris
                'header_text': RGBColor(255, 255, 255),
                'border_color': RGBColor(127, 140, 141),
                'alt_row_bg': None  # Pas d'alternance
            }
        }
    
    def _shade_cell(self, cell, color: RGBColor):
        """
        Applique une couleur de fond à une cellule (méthode compatible)
        
        Args:
            cell: Cellule Word
            color: Couleur RGB
        """
        try:
            # Méthode 1 : Essayer avec shading (nouvelles versions)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color.rgb_bytes.hex())
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception:
            try:
                # Méthode 2 : Approche alternative
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:shd')
                tcVAlign.set(qn('w:fill'), color.rgb_bytes.hex())
                tcPr.append(tcVAlign)
            except Exception:
                # Si tout échoue, on ignore le shading
                pass
    
    def format_value(self, value: Any, precision: int = 2) -> str:
        """
        Formate une valeur pour affichage
        
        Args:
            value: Valeur à formater
            precision: Nombre de décimales pour les floats
            
        Returns:
            Valeur formatée en string
        """
        if pd.isna(value):
            return "N/A"
        elif isinstance(value, float):
            return f"{value:.{precision}f}"
        elif isinstance(value, int):
            return str(value)
        else:
            return str(value)
    
    def dataframe_to_docx_table(
        self,
        doc: Document,
        df: pd.DataFrame,
        title: Optional[str] = None,
        precision: int = 2,
        include_index: bool = False
    ) -> Document:
        """
        Ajoute un DataFrame formaté à un document Word
        
        Args:
            doc: Document Word (python-docx)
            df: DataFrame à ajouter
            title: Titre optionnel du tableau
            precision: Nombre de décimales
            include_index: Inclure l'index du DataFrame
            
        Returns:
            Document Word modifié
        """
        # Ajouter le titre si fourni
        if title:
            heading = doc.add_heading(title, level=3)
            heading.style.font.color.rgb = RGBColor(44, 62, 80)
        
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Créer le tableau
        num_rows = len(df_display) + 1  # +1 pour le header
        num_cols = len(df_display.columns)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Configuration du style
        style_cfg = self.styles_config.get(self.style, self.styles_config['professional'])
        
        # HEADER ROW
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df_display.columns):
            cell = header_cells[i]
            cell.text = str(col_name)
            
            # Style du header
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = style_cfg['header_text']
            
            # Couleur de fond du header (compatible)
            self._shade_cell(cell, style_cfg['header_bg'])
            
            # Alignement
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # DATA ROWS
        for row_idx, row_data in enumerate(df_display.itertuples(index=False), start=1):
            row_cells = table.rows[row_idx].cells
            
            for col_idx, value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = self.format_value(value, precision)
                
                # Style de la cellule
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
                run.font.size = Pt(10)
                
                # Alignement : nombres à droite, texte à gauche
                if isinstance(value, (int, float)) and not pd.isna(value):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Alternance des couleurs de lignes
                if style_cfg['alt_row_bg'] and row_idx % 2 == 0:
                    self._shade_cell(cell, style_cfg['alt_row_bg'])
        
        # Ajouter un espace après le tableau
        doc.add_paragraph()
        
        return doc
    
    def dataframe_to_html(
        self,
        df: pd.DataFrame,
        title: Optional[str] = None,
        precision: int = 2,
        include_index: bool = False
    ) -> str:
        """
        Convertit un DataFrame en HTML professionnel
        
        Args:
            df: DataFrame à convertir
            title: Titre optionnel
            precision: Nombre de décimales
            include_index: Inclure l'index
            
        Returns:
            String HTML
        """
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Formater les valeurs
        df_formatted = df_display.copy()
        for col in df_formatted.columns:
            df_formatted[col] = df_formatted[col].apply(
                lambda x: self.format_value(x, precision)
            )
        
        # Configuration du style
        style_cfg = self.styles_config.get(self.style, self.styles_config['professional'])
        header_bg = f"rgb{style_cfg['header_bg']}"
        header_text = f"rgb{style_cfg['header_text']}"
        alt_row_bg = f"rgb{style_cfg['alt_row_bg']}" if style_cfg['alt_row_bg'] else "transparent"
        
        # Générer le HTML
        html = ""
        
        if title:
            html += f'<h3 style="color: #2c3e50; margin-top: 20px;">{title}</h3>\n'
        
        html += f'''
<table style="border-collapse: collapse; width: 100%; margin: 20px 0; font-family: Arial, sans-serif;">
    <thead>
        <tr style="background-color: {header_bg}; color: {header_text};">
'''
        
        # Headers
        for col in df_formatted.columns:
            html += f'            <th style="border: 1px solid #ddd; padding: 12px; text-align: left; font-weight: bold;">{col}</th>\n'
        
        html += '        </tr>\n    </thead>\n    <tbody>\n'
        
        # Rows
        for idx, row in df_formatted.iterrows():
            bg_color = alt_row_bg if idx % 2 == 0 else "white"
            html += f'        <tr style="background-color: {bg_color};">\n'
            
            for value in row:
                html += f'            <td style="border: 1px solid #ddd; padding: 10px;">{value}</td>\n'
            
            html += '        </tr>\n'
        
        html += '    </tbody>\n</table>\n'
        
        return html
    
    def dataframe_to_markdown(
        self,
        df: pd.DataFrame,
        precision: int = 2,
        include_index: bool = False
    ) -> str:
        """
        Convertit un DataFrame en Markdown
        
        Args:
            df: DataFrame à convertir
            precision: Nombre de décimales
            include_index: Inclure l'index
            
        Returns:
            String Markdown
        """
        # Préparer les données
        if include_index:
            df_display = df.reset_index()
        else:
            df_display = df.copy()
        
        # Formater les valeurs
        df_formatted = df_display.copy()
        for col in df_formatted.columns:
            df_formatted[col] = df_formatted[col].apply(
                lambda x: self.format_value(x, precision)
            )
        
        # Générer le Markdown
        md = "| " + " | ".join(str(col) for col in df_formatted.columns) + " |\n"
        md += "|" + "|".join(["---" for _ in df_formatted.columns]) + "|\n"
        
        for _, row in df_formatted.iterrows():
            md += "| " + " | ".join(str(val) for val in row) + " |\n"
        
        return md


# ============= FONCTIONS UTILITAIRES =============

def format_statistics_table(df: pd.DataFrame, column: str, style: str = "professional") -> Dict[str, Any]:
    """
    Crée un tableau de statistiques descriptives formaté
    
    Args:
        df: DataFrame source
        column: Nom de la colonne à analyser
        style: Style du tableau
        
    Returns:
        Dictionnaire avec les différents formats
    """
    formatter = TableFormatter(style=style)
    
    # Calculer les statistiques
    stats = df[column].describe()
    stats_df = pd.DataFrame({
        'Statistique': stats.index,
        'Valeur': stats.values
    })
    
    return {
        'dataframe': stats_df,
        'html': formatter.dataframe_to_html(stats_df, title=f"Statistiques descriptives : {column}"),
        'markdown': formatter.dataframe_to_markdown(stats_df)
    }


def format_crosstab(df: pd.DataFrame, row_var: str, col_var: str, style: str = "professional") -> Dict[str, Any]:
    """
    Crée un tableau croisé formaté
    
    Args:
        df: DataFrame source
        row_var: Variable en lignes
        col_var: Variable en colonnes
        style: Style du tableau
        
    Returns:
        Dictionnaire avec les différents formats
    """
    formatter = TableFormatter(style=style)
    
    # Créer le tableau croisé
    crosstab = pd.crosstab(df[row_var], df[col_var])
    
    return {
        'dataframe': crosstab,
        'html': formatter.dataframe_to_html(
            crosstab, 
            title=f"Tableau croisé : {row_var} × {col_var}",
            include_index=True
        ),
        'markdown': formatter.dataframe_to_markdown(crosstab, include_index=True)
    }


def format_correlation_matrix(df: pd.DataFrame, style: str = "professional") -> Dict[str, Any]:
    """
    Crée une matrice de corrélation formatée
    
    Args:
        df: DataFrame source (colonnes numériques)
        style: Style du tableau
        
    Returns:
        Dictionnaire avec les différents formats
    """
    formatter = TableFormatter(style=style)
    
    # Calculer la matrice de corrélation
    corr_matrix = df.select_dtypes(include=['number']).corr()
    
    return {
        'dataframe': corr_matrix,
        'html': formatter.dataframe_to_html(
            corr_matrix,
            title="Matrice de corrélation",
            precision=3,
            include_index=True
        ),
        'markdown': formatter.dataframe_to_markdown(corr_matrix, precision=3, include_index=True)
    }


# ============= EXEMPLE D'UTILISATION =============

if __name__ == "__main__":
    # Exemple avec des données fictives
    import numpy as np
    
    # Créer un DataFrame de test
    np.random.seed(42)
    test_df = pd.DataFrame({
        'Variable': ['Âge', 'Salaire', 'Expérience', 'Satisfaction'],
        'Moyenne': np.random.uniform(20, 60, 4),
        'Écart-type': np.random.uniform(5, 15, 4),
        'Min': np.random.uniform(10, 30, 4),
        'Max': np.random.uniform(70, 100, 4)
    })
    
    # Tester les différents formats
    formatter = TableFormatter(style="professional")
    
    print("=== HTML ===")
    print(formatter.dataframe_to_html(test_df, title="Statistiques descriptives"))
    
    print("\n=== MARKDOWN ===")
    print(formatter.dataframe_to_markdown(test_df))
    
    print("\n=== WORD ===")
    doc = Document()
    doc = formatter.dataframe_to_docx_table(doc, test_df, title="Statistiques descriptives")
    
    # Sauvegarder
    buffer = io.BytesIO()
    doc.save(buffer)
    print(f"Document Word créé ({len(buffer.getvalue())} bytes)")