"""
AI Statistical Reporter - Application Streamlit
Version professionnelle avec export et analyse de données
Version améliorée avec gestion d'erreurs, cache, rate limiting, autosave
"""

import streamlit as st
import pandas as pd
import json
import os
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
import io
import base64

# ═══════════════════════════════════════════════════════════════
# INITIALISATION SESSION_STATE (CRITIQUE - EN PREMIER)
# ═══════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════
# INITIALISATION SESSION_STATE (CRITIQUE - EN PREMIER)
# ═══════════════════════════════════════════════════════════════

# Initialiser TOUTES les variables AVANT toute utilisation
if 'writing_profile' not in st.session_state:
    st.session_state.writing_profile = None

if 'use_context' not in st.session_state:
    st.session_state.use_context = None  # [OK] IMPORTANT

if 'study_context' not in st.session_state:
    st.session_state.study_context = None

if 'analysis_mode' not in st.session_state:
    st.session_state.analysis_mode = None

if 'temp_path' not in st.session_state:
    st.session_state['temp_path'] = None  # [OK] IMPORTANT

if 'csv_data' not in st.session_state:
    st.session_state.csv_data = None

if 'plan' not in st.session_state:
    st.session_state.plan = None

if 'workflow_step' not in st.session_state:
    st.session_state.workflow_step = 1

if 'workflow_history' not in st.session_state:
    st.session_state.workflow_history = []
# Initialiser TOUTES les variables AVANT toute utilisation
if 'writing_profile' not in st.session_state:
    st.session_state.writing_profile = None

if 'use_context' not in st.session_state:
    st.session_state.use_context = None

if 'study_context' not in st.session_state:
    st.session_state.study_context = None

if 'analysis_mode' not in st.session_state:
    st.session_state.analysis_mode = None

# ═══════════════════════════════════════════════════════════════

# CSS personnalisé pour améliorer l'UX
st.markdown("""
<style>
    /* Bouton Enregistrer (vert) */
    div.stButton > button[kind="primary"] {
        background-color: #28a745;
        color: white;
        border: none;
    }
    div.stButton > button[kind="primary"]:hover {
        background-color: #218838;
        color: white;
    }
    
    /* Bouton Annuler (rouge) */
    div.stButton > button:not([kind="primary"]):not([disabled]) {
        background-color: #dc3545;
        color: white;
        border: none;
    }
    div.stButton > button:not([kind="primary"]):not([disabled]):hover {
        background-color: #c82333;
        color: white;
    }
    
    /* Bouton désactivé (gris) */
    div.stButton > button[disabled] {
        background-color: #6c757d;
        color: #ffffff80;
        cursor: not-allowed;
    }
    
    /* Zone de texte d'édition */
    textarea {
        font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
        font-size: 14px;
        line-height: 1.6;
    }
    
    /* Messages de succès/erreur plus visibles */
    .stSuccess {
        padding: 1rem;
        border-radius: 0.5rem;
    }
    
    .stWarning {
        padding: 1rem;
        border-radius: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# Importer les modules existants
from week2_architect_agent import analyze_csv, generate_report_plan
from table_formatter import TableFormatter


# Workflow Manager
try:
    from workflow_manager import WorkflowManager, WorkflowStep
    WORKFLOW_MANAGER_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] workflow_manager.py non disponible: {e}")
    WORKFLOW_MANAGER_AVAILABLE = False

# ===== NOUVEAUX MODULES =====
# Gestion d'erreurs
try:
    from error_handler import (
        handle_errors,
        validate_api_keys,
        validate_csv_file,
        APIError,
        DataError,
        ParsingError,
        safe_execute
    )
    ERROR_HANDLER_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] error_handler.py non disponible: {e}")
    ERROR_HANDLER_AVAILABLE = False
except Exception as e:
    print(f"[WARNING] Erreur dans error_handler: {e}")
    ERROR_HANDLER_AVAILABLE = False

# Logging
try:
    from logging_config import setup_logging, log_api_call, log_user_action
    logger = setup_logging()
    logger.info("Application démarrée")
    LOGGING_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] logging_config.py non disponible: {e}")
    LOGGING_AVAILABLE = False
    import logging
    logger = logging.getLogger(__name__)
except Exception as e:
    print(f"[WARNING] Erreur lors du setup logging: {e}")
    LOGGING_AVAILABLE = False
    import logging
    logger = logging.getLogger(__name__)

# Cache
try:
    from cache_manager import cached_plan_generation, cached_data_analysis, display_cache_info, cache
    CACHE_AVAILABLE = True
except ImportError:
    print("[WARNING] cache_manager.py non disponible")
    CACHE_AVAILABLE = False

# Rate limiting
try:
    from rate_limiter import rate_limiter, RATE_LIMITS, display_rate_limit_info
    RATE_LIMIT_AVAILABLE = True
except ImportError:
    print("[WARNING] rate_limiter.py non disponible")
    RATE_LIMIT_AVAILABLE = False

# Autosave
try:
    from autosave import autosave, enable_autosave_for_plan, show_draft_recovery, show_draft_manager
    AUTOSAVE_AVAILABLE = True
except ImportError:
    print("[WARNING] autosave.py non disponible")
    AUTOSAVE_AVAILABLE = False

# UI Components
try:
    from ui_components import (
        show_progress_steps,
        show_success_message,
        show_file_upload_zone,
        show_onboarding_tour
    )
    UI_COMPONENTS_AVAILABLE = True
except ImportError:
    print("[WARNING] ui_components.py non disponible")
    UI_COMPONENTS_AVAILABLE = False

# ===== MODULES SEMAINE 6 : PERSISTANCE & MÉMOIRE =====
# E2B Session Manager
try:
    from e2b_session_manager import (
        get_sandbox_for_user,
        execute_python_code,
        display_session_status_in_streamlit,
        session_manager
    )
    E2B_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] e2b_session_manager.py non disponible: {e}")
    E2B_AVAILABLE = False

# Contextual Memory
try:
    from contextual_memory import (
        add_chapter_to_memory,
        get_context_for_chapter,
        display_memory_in_streamlit,
        contextual_memory
    )
    MEMORY_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] contextual_memory.py non disponible: {e}")
    MEMORY_AVAILABLE = False

# Chapter Workflow
try:
    from chapter_workflow import (
        initialize_workflow,
        display_workflow_progress,
        ReportGenerationWorkflow
    )
    WORKFLOW_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] chapter_workflow.py non disponible: {e}")
    WORKFLOW_AVAILABLE = False

# Translations / Multilingue
try:
    from translations import get_text as t, get_language_name
    TRANSLATIONS_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] translations.py non disponible: {e}")
    TRANSLATIONS_AVAILABLE = False
    # Fallback : fonction qui retourne le texte tel quel
    def t(key, lang='fr', **kwargs):
        return key
    def get_language_name(lang):
        return lang

# Cost Controller
try:
    from cost_controller import (
        cost_controller,
        get_length_guidelines,
        display_cost_summary_in_streamlit,
        ChapterConfig
    )
    COST_CONTROLLER_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] cost_controller.py non disponible: {e}")
    COST_CONTROLLER_AVAILABLE = False

# Study Context
try:
    from study_context import StudyContext, study_context
    STUDY_CONTEXT_AVAILABLE = True
except ImportError as e:
    print(f"[WARNING] study_context.py non disponible: {e}")
    STUDY_CONTEXT_AVAILABLE = False

load_dotenv()

# Validation des clés API au démarrage
if ERROR_HANDLER_AVAILABLE:
    try:
        validate_api_keys()
    except APIError as e:
        st.error(e.user_message)
        st.stop()

# Configuration de la page
st.set_page_config(
    page_title="AI Statistical Reporter",
    page_icon="[DATA]",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé professionnel
st.markdown("""
<style>
    /* Layout général */
    .main {
        padding: 0rem 1rem;
    }
    
    /* Boutons */
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    /* Zone de texte */
    .stTextArea>div>div>textarea {
        font-family: 'Courier New', monospace;
        font-size: 14px;
        line-height: 1.6;
        border-radius: 5px;
    }
    
    /* Messages de statut */
    .success-box {
        padding: 1rem;
        border-radius: 5px;
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        color: #155724;
        margin: 1rem 0;
    }
    
    .info-box {
        padding: 1rem;
        border-radius: 5px;
        background-color: #d1ecf1;
        border-left: 4px solid #17a2b8;
        color: #0c5460;
        margin: 1rem 0;
    }
    
    .warning-box {
        padding: 1rem;
        border-radius: 5px;
        background-color: #fff3cd;
        border-left: 4px solid #ffc107;
        color: #856404;
        margin: 1rem 0;
    }
    
    .error-box {
        padding: 1rem;
        border-radius: 5px;
        background-color: #f8d7da;
        border-left: 4px solid #dc3545;
        color: #721c24;
        margin: 1rem 0;
    }
    
    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #f8f9fa;
    }
    
    /* Titres */
    h1 {
        color: #2c3e50;
        font-weight: 700;
    }
    
    h2 {
        color: #34495e;
        font-weight: 600;
    }
    
    h3 {
        color: #7f8c8d;
        font-weight: 600;
    }
    
    /* Métrics */
    div[data-testid="stMetricValue"] {
        font-size: 2rem;
        font-weight: 700;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        font-weight: 600;
        background-color: #f8f9fa;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)


# Fonctions utilitaires
def json_to_editable_text(plan: dict) -> str:
    """Convertit le JSON du plan en format texte éditable"""
    text = f"TITRE: {plan.get('titre', 'Rapport Statistique')}\n"
    text += f"DATE: {plan.get('date', datetime.now().strftime('%Y-%m-%d'))}\n"
    text += f"AUTEUR: {plan.get('auteur', 'AI Reporter')}\n\n"
    
    for i, chap in enumerate(plan.get('chapitres', []), 1):
        text += f"{i}. {chap.get('titre', f'Chapitre {i}')}\n"
        
        for j, sec in enumerate(chap.get('sections', []), 1):
            text += f"   {i}.{j}. {sec.get('titre', f'Section {j}')}\n"
            
            for analyse in sec.get('analyses', []):
                text += f"      - {analyse}\n"
            
            text += "\n"
    
    return text


def text_to_json_with_ai(text: str) -> dict:
    """Parse le texte modifié et génère un JSON valide avec l'IA"""
    
    from google.genai import Client as GminiClient
    from langchain_anthropic import ChatAnthropic
    
    # AGENT PARSER
    parse_prompt = f"""
Tu es un agent spécialisé dans le parsing de plans de rapports.

TEXTE DU PLAN:
```
{text}
```

TÂCHE:
Convertis ce texte en JSON avec ce format EXACT:
{{
  "titre": "...",
  "date": "...",
  "auteur": "...",
  "chapitres": [
    {{
      "numero": "1",
      "titre": "...",
      "sections": [
        {{
          "titre": "...",
          "analyses": ["...", "..."]
        }}
      ]
    }}
  ]
}}

RÈGLES:
1. "TITRE:", "DATE:", "AUTEUR:" → métadonnées
2. "1.", "2." → chapitres
3. "   1.1.", "   1.2." → sections (indentées)
4. "      - " → analyses
5. Si une section n'a pas d'analyses, invente-en 2-3 pertinentes
6. Garde EXACTEMENT les titres de l'utilisateur

Retourne UNIQUEMENT le JSON, sans markdown.
"""
    
    # Essayer Gemini
    gmini_key = os.getenv("GMINI_API_KEY")
    parsed_json = None
    
    if gmini_key:
        try:
            gclient = GminiClient(api_key=gmini_key)
            chat = gclient.chats.create(model="gemini-2.5-flash")
            gres = chat.send_message(parse_prompt)
            
            gen = None
            if hasattr(gres, "candidates") and gres.candidates:
                first = gres.candidates[0]
                if hasattr(first, "content"):
                    gen = first.content
                    if not isinstance(gen, str) and hasattr(gen, "parts"):
                        parts = getattr(gen, "parts") or []
                        texts = [getattr(p, "text", "") for p in parts if getattr(p, "text", None)]
                        gen = "\n".join(texts).strip()
            
            if gen:
                parsed_json = str(gen)
        except Exception as e:
            st.error(f"Erreur Gemini: {e}")
    
    # Fallback Anthropic
    if not parsed_json:
        try:
            llm = ChatAnthropic(
                model="claude-sonnet-4-20250514",
                api_key=os.getenv("ANTHROPIC_API_KEY"),
                temperature=0.1
            )
            response = llm.invoke(parse_prompt)
            parsed_json = response.content
        except Exception as e:
            raise Exception(f"Erreur parsing: {e}")
    
    # Nettoyer
    if "```json" in parsed_json:
        parsed_json = parsed_json.split("```json")[1].split("```")[0].strip()
    elif "```" in parsed_json:
        parsed_json = parsed_json.split("```")[1].split("```")[0].strip()
    
    return json.loads(parsed_json)



def regenerate_plan_with_instructions(current_plan: dict, instructions: str, metadata: dict, 
                                     keep_structure: bool = False, academic: bool = False, 
                                     detailed: bool = False, study_context=None) -> dict:
    """
    Régénère un plan en prenant en compte les instructions de l'utilisateur
    
    Args:
        current_plan: Plan actuel (dict JSON)
        instructions: Instructions de modification (texte libre)
        metadata: Métadonnées du CSV
        keep_structure: Conserver la structure actuelle
        academic: Style académique
        detailed: Mode détaillé
        study_context: Contexte de l'étude (optionnel)
    
    Returns:
        Nouveau plan (dict JSON)
    """
    import os
    import json
    
    # Convertir le plan actuel en texte lisible
    current_plan_text = json_to_editable_text(current_plan)
    
    # Construire le prompt de régénération
    prompt = f"""
Tu es un assistant IA spécialisé dans la création de plans de rapports statistiques.

════════════════════════════════════════════════════════════════
PLAN ACTUEL
════════════════════════════════════════════════════════════════

{current_plan_text}

════════════════════════════════════════════════════════════════
INSTRUCTIONS DE L'UTILISATEUR
════════════════════════════════════════════════════════════════

{instructions}

════════════════════════════════════════════════════════════════
CONTRAINTES
════════════════════════════════════════════════════════════════

- {"CONSERVER la structure actuelle (nombre et ordre des chapitres)" if keep_structure else "Tu peux MODIFIER la structure librement"}
- {"Style ACADÉMIQUE (formel, références, méthodologie)" if academic else "Style ACCESSIBLE (vulgarisé, pédagogique)"}
- {"Mode DÉTAILLÉ (plus de sections et analyses par chapitre)" if detailed else "Mode STANDARD"}

════════════════════════════════════════════════════════════════
DONNÉES DISPONIBLES
════════════════════════════════════════════════════════════════

Variables : {', '.join(metadata.get('colonnes', []))}
Nombre de lignes : {metadata.get('nombre_lignes', 'N/A')}
Types de variables :
{chr(10).join(f"  - {col}: {dtype}" for col, dtype in metadata.get('types', {}).items())}

"""
    
    # Ajouter le contexte d'étude si disponible
    if study_context:
        try:
            prompt += f"""
════════════════════════════════════════════════════════════════
CONTEXTE DE L'ÉTUDE
════════════════════════════════════════════════════════════════

{study_context.to_prompt_context()}

"""
        except:
            pass
    
    prompt += """
════════════════════════════════════════════════════════════════
TÂCHE
════════════════════════════════════════════════════════════════

Génère un NOUVEAU PLAN qui :
1. Prend en compte TOUTES les instructions de l'utilisateur
2. Respecte les contraintes définies
3. Utilise UNIQUEMENT les variables disponibles dans les données
4. Est structuré en chapitres, sections et analyses détaillées

FORMAT DE SORTIE (JSON) :

{
  "titre": "Titre du rapport",
  "date": "2025-12-26",
  "auteur": "AI Statistical Reporter",
  "chapitres": [
    {
      "numero": "1",
      "titre": "Titre du chapitre",
      "sections": [
        {
          "titre": "Titre de la section",
          "analyses": [
            "Analyse 1",
            "Analyse 2"
          ]
        }
      ]
    }
  ]
}

Retourne UNIQUEMENT le JSON, sans texte avant ou après.
"""
    
    # Appeler l'IA (Gemini ou Claude)
    try:
        # Essayer Gemini d'abord
        import google.generativeai as genai
        
        genai.configure(api_key=os.getenv("GMINI_API_KEY"))
        model_name = os.getenv("GEMINI_MODEL_PLAN", "gemini-2.5-flash")
        model = genai.GenerativeModel(model_name)
        #model = genai.GenerativeModel("gemini-2.0-flash-exp")
        
        response = model.generate_content(prompt)
        result_text = response.text
        
    except Exception as e:
        # Fallback Claude
        try:
            from anthropic import Anthropic
            
            client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
            
            response = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            result_text = response.content[0].text
        
        except Exception as e2:
            raise Exception(f"Échec Gemini et Claude. Gemini: {e}, Claude: {e2}")
    
    # Nettoyer le JSON
    if "```json" in result_text:
        result_text = result_text.split("```json")[1].split("```")[0].strip()
    elif "```" in result_text:
        result_text = result_text.split("```")[1].split("```")[0].strip()
    
    # Parser le JSON
    new_plan = json.loads(result_text)
    
    return new_plan

def display_plan_preview(plan: dict):
    """Affiche un aperçu formaté du plan"""
    
    # Option d'affichage
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown(f"### {plan.get('titre', 'Plan du Rapport')}")
        st.caption(f"Date: {plan.get('date', 'Aujourdhui')} | Auteur: {plan.get('auteur', 'AI Reporter')}")
    
    with col2:
        show_all = st.checkbox("Afficher tout le contenu", value=True, key="show_all_preview")
    
    total_sections = 0
    total_analyses = 0
    
    for i, chap in enumerate(plan.get('chapitres', []), 1):
        with st.expander(f"{i}. {chap.get('titre')}", expanded=False):
            for j, sec in enumerate(chap.get('sections', []), 1):
                st.markdown(f"**{i}.{j}. {sec.get('titre')}**")
                analyses = sec.get('analyses', [])
                total_sections += 1
                total_analyses += len(analyses)
                
                if show_all:
                    for analyse in analyses:
                        st.markdown(f"   - {analyse}")
                else:
                    for analyse in analyses[:3]:
                        st.markdown(f"   - {analyse}")
                    
                    if len(analyses) > 3:
                        st.markdown(f"   - *... et {len(analyses) - 3} autres analyses*")
                
                st.markdown("")
    
    # Statistiques
    st.divider()
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Chapitres", len(plan.get('chapitres', [])))
    with col2:
        st.metric("Sections", total_sections)
    with col3:
        st.metric("Analyses", total_analyses)


def save_plan_to_file(plan: dict, filename: str = "report_plan.json"):
    """Sauvegarde le plan en JSON"""
    output_dir = Path("output/plans")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = output_dir / filename
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(plan, f, indent=2, ensure_ascii=False)
    
    return filepath


def plan_to_html(plan: dict) -> str:
    """Convertit le plan en HTML professionnel"""
    
    html = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{plan.get('titre', 'Plan de Rapport')}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            border-left: 4px solid #3498db;
            padding-left: 15px;
        }}
        h3 {{
            color: #7f8c8d;
            margin-top: 20px;
            margin-left: 20px;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .metadata p {{
            margin: 5px 0;
        }}
        ul {{
            margin-left: 40px;
        }}
        li {{
            margin: 8px 0;
        }}
        .stats {{
            background-color: #e8f5e9;
            padding: 15px;
            border-radius: 5px;
            margin-top: 30px;
            border-left: 4px solid #4caf50;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 20px;
            margin-top: 15px;
        }}
        .stat-item {{
            text-align: center;
        }}
        .stat-value {{
            font-size: 2em;
            font-weight: bold;
            color: #2c3e50;
        }}
        .stat-label {{
            color: #7f8c8d;
            font-size: 0.9em;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ecf0f1;
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <h1>{plan.get('titre', 'Plan de Rapport')}</h1>
    
    <div >
        <p><strong>Date :</strong> {plan.get('date', 'Non spécifiée')}</p>
        <p><strong>Auteur :</strong> {plan.get('auteur', 'Non spécifié')}</p>
    </div>
"""
    
    # Compter les éléments
    total_sections = 0
    total_analyses = 0
    
    # Générer le contenu
    for i, chap in enumerate(plan.get('chapitres', []), 1):
        html += f"    <h2>{i}. {chap.get('titre', f'Chapitre {i}')}</h2>\n"
        
        for j, sec in enumerate(chap.get('sections', []), 1):
            total_sections += 1
            html += f"    <h3>{i}.{j}. {sec.get('titre', f'Section {j}')}</h3>\n"
            html += "    <ul>\n"
            
            analyses = sec.get('analyses', [])
            total_analyses += len(analyses)
            
            for analyse in analyses:
                html += f"        <li>{analyse}</li>\n"
            
            html += "    </ul>\n"
    
    # Ajouter les statistiques
    html += f"""
    <div >
        <h2>Statistiques du plan</h2>
        <div >
            <div >
                <div >{len(plan.get('chapitres', []))}</div>
                <div >Chapitres</div>
            </div>
            <div >
                <div >{total_sections}</div>
                <div >Sections</div>
            </div>
            <div >
                <div >{total_analyses}</div>
                <div >Analyses prévues</div>
            </div>
        </div>
    </div>
    
    <div >
        <p>Généré par AI Statistical Reporter - {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </div>
</body>
</html>
"""
    
    return html


def plan_to_docx(plan: dict) -> io.BytesIO:
    """Convertit le plan en DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx non installé. Installez avec: pip install python-docx")
    
    doc = Document()
    
    # Titre
    title = doc.add_heading(plan.get('titre', 'Plan de Rapport'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Métadonnées
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(f"Date : {plan.get('date', 'Non spécifiée')}\n").bold = True
    meta.add_run(f"Auteur : {plan.get('auteur', 'Non spécifié')}").bold = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Chapitres
    for i, chap in enumerate(plan.get('chapitres', []), 1):
        # Titre du chapitre
        doc.add_heading(f"{i}. {chap.get('titre', f'Chapitre {i}')}", 1)
        
        # Sections
        for j, sec in enumerate(chap.get('sections', []), 1):
            doc.add_heading(f"{i}.{j}. {sec.get('titre', f'Section {j}')}", 2)
            
            # Analyses
            for analyse in sec.get('analyses', []):
                p = doc.add_paragraph(analyse, style='List Bullet')
        
        doc.add_paragraph()
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def plan_to_pdf(plan: dict) -> bytes:
    """Convertit le plan en PDF via HTML"""
    try:
        from weasyprint import HTML
        html_content = plan_to_html(plan)
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes
    except ImportError as e:
        raise ImportError(f"WeasyPrint non installé ou bibliothèques système manquantes: {e}")
    except Exception as e:
        raise Exception(f"Erreur lors de la génération du PDF: {e}")


def analyze_data_quality(df: pd.DataFrame) -> dict:
    """Analyse détaillée de la qualité des données"""
    
    analysis = {
        'shape': df.shape,
        'columns': list(df.columns),
        'dtypes': df.dtypes.to_dict(),
        'numeric_cols': list(df.select_dtypes(include=['number']).columns),
        'categorical_cols': list(df.select_dtypes(include=['object', 'category']).columns),
        'missing_values': {},
        'numeric_stats': {},
        'categorical_stats': {}
    }
    
    # Valeurs manquantes
    for col in df.columns:
        missing_count = df[col].isnull().sum()
        missing_pct = (missing_count / len(df)) * 100
        analysis['missing_values'][col] = {
            'count': int(missing_count),
            'percentage': round(missing_pct, 2)
        }
    
    # Statistiques numériques
    for col in analysis['numeric_cols']:
        analysis['numeric_stats'][col] = {
            'count': int(df[col].count()),
            'mean': round(df[col].mean(), 2) if df[col].count() > 0 else None,
            'std': round(df[col].std(), 2) if df[col].count() > 0 else None,
            'min': round(df[col].min(), 2) if df[col].count() > 0 else None,
            'q25': round(df[col].quantile(0.25), 2) if df[col].count() > 0 else None,
            'median': round(df[col].median(), 2) if df[col].count() > 0 else None,
            'q75': round(df[col].quantile(0.75), 2) if df[col].count() > 0 else None,
            'max': round(df[col].max(), 2) if df[col].count() > 0 else None
        }
    
    # Statistiques catégorielles
    for col in analysis['categorical_cols']:
        value_counts = df[col].value_counts()
        analysis['categorical_stats'][col] = {
            'unique_values': int(df[col].nunique()),
            'most_common': value_counts.head(10).to_dict(),
            'count': int(df[col].count())
        }
    
    return analysis


def display_data_overview(df: pd.DataFrame):
    """Affiche un aperçu détaillé des données"""
    
    st.markdown("### Aperçu général des données")
    
    # Dimensions
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Nombre de lignes", f"{len(df):,}")
    with col2:
        st.metric("Nombre de colonnes", len(df.columns))
    with col3:
        missing_total = df.isnull().sum().sum()
        missing_pct = (missing_total / (len(df) * len(df.columns))) * 100
        st.metric("Valeurs manquantes", f"{missing_pct:.1f}%")
    
    st.markdown("---")
    
    # Onglets pour différentes vues
    tab1, tab2, tab3, tab4 = st.tabs([
        "Aperçu des données",
        "Variables quantitatives", 
        "Variables qualitatives",
        "Valeurs manquantes"
    ])
    
    with tab1:
        st.markdown("#### Premières lignes du dataset")
        st.dataframe(df.head(10), use_container_width=True)
        
        st.markdown("#### Types de variables")
        type_df = pd.DataFrame({
            'Variable': df.columns,
            'Type': df.dtypes.astype(str),
            'Valeurs non-nulles': df.count().values,
            'Valeurs uniques': [df[col].nunique() for col in df.columns]
        })
        st.dataframe(type_df, use_container_width=True)
    
    with tab2:
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        
        if numeric_cols:
            st.markdown(f"#### Variables quantitatives ({len(numeric_cols)})")
            
            for col in numeric_cols:
                with st.expander(f"{col}"):
                    col_stats = df[col].describe()
                    
                    # [OK] CORRECTION : Tableau formaté
                    formatter = TableFormatter(style='professional')
                    stats_df = pd.DataFrame({
                        'Statistique': col_stats.index,
                        'Valeur': col_stats.values
                    })
                    stats_html = formatter.dataframe_to_html(
                        stats_df,
                        title=f"Statistiques descriptives : {col}",
                        precision=2
                    )
                    st.markdown(stats_html, unsafe_allow_html=True)
                    st.markdown("---")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Moyenne", f"{col_stats['mean']:.2f}")
                        st.metric("Écart-type", f"{col_stats['std']:.2f}")
                    with col2:
                        st.metric("Minimum", f"{col_stats['min']:.2f}")
                        st.metric("Q1 (25%)", f"{col_stats['25%']:.2f}")
                    with col3:
                        st.metric("Médiane", f"{col_stats['50%']:.2f}")
                        st.metric("Q3 (75%)", f"{col_stats['75%']:.2f}")
                    with col4:
                        st.metric("Maximum", f"{col_stats['max']:.2f}")
                        st.metric("Count", f"{int(col_stats['count'])}")
                    # Histogramme
                    st.markdown("**Distribution**")
                    try:
                        import matplotlib.pyplot as plt
                        fig, ax = plt.subplots(figsize=(8, 3))
                        df[col].hist(bins=20, ax=ax, edgecolor='black')
                        ax.set_xlabel(col)
                        ax.set_ylabel('Fréquence')
                        ax.set_title(f'Distribution de {col}')
                        st.pyplot(fig)
                        plt.close()
                    except:
                        st.info("Matplotlib non disponible pour l'affichage des graphiques")
        else:
            st.info("Aucune variable quantitative détectée")
    
    with tab3:
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        if categorical_cols:
            st.markdown(f"#### Variables qualitatives ({len(categorical_cols)})")
            
            for col in categorical_cols:
                with st.expander(f"{col}"):
                    unique_count = df[col].nunique()
                    total_count = df[col].count()
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Modalités uniques", unique_count)
                    with col2:
                        st.metric("Valeurs non-nulles", total_count)
                    
                    st.markdown("**Répartition des modalités**")
                    
                    value_counts = df[col].value_counts()
                    value_pcts = df[col].value_counts(normalize=True) * 100
                    
                    freq_df = pd.DataFrame({
                        'Modalité': value_counts.index[:20],
                        'Effectif': value_counts.values[:20],
                        'Pourcentage': [f"{pct:.1f}%" for pct in value_pcts.values[:20]]
                    })
                    
                    st.dataframe(freq_df, use_container_width=True)
                    
                    if unique_count > 20:
                        st.caption(f"Affichage des 20 modalités les plus fréquentes sur {unique_count}")
                    
                    # Graphique en barres
                    try:
                        import matplotlib.pyplot as plt
                        fig, ax = plt.subplots(figsize=(8, 4))
                        value_counts.head(10).plot(kind='bar', ax=ax)
                        ax.set_xlabel('Modalités')
                        ax.set_ylabel('Fréquence')
                        ax.set_title(f'Top 10 des modalités - {col}')
                        plt.xticks(rotation=45, ha='right')
                        plt.tight_layout()
                        st.pyplot(fig)
                        plt.close()
                    except:
                        pass
        else:
            st.info("Aucune variable qualitative détectée")
    
    with tab4:
        st.markdown("#### Analyse des valeurs manquantes")
        
        missing_data = []
        for col in df.columns:
            missing_count = df[col].isnull().sum()
            missing_pct = (missing_count / len(df)) * 100
            missing_data.append({
                'Variable': col,
                'Valeurs manquantes': missing_count,
                'Pourcentage': f"{missing_pct:.2f}%",
                'Valeurs présentes': len(df) - missing_count
            })
        
        missing_df = pd.DataFrame(missing_data)
        missing_df = missing_df.sort_values('Valeurs manquantes', ascending=False)
        
        # Filtrer pour afficher seulement celles avec des manquantes
        missing_df_filtered = missing_df[missing_df['Valeurs manquantes'] > 0]
        
        if len(missing_df_filtered) > 0:
            st.dataframe(missing_df_filtered, use_container_width=True)
            
            # Graphique des manquantes
            try:
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(10, 6))
                missing_df_filtered.plot(
                    x='Variable', 
                    y='Valeurs manquantes',
                    kind='barh',
                    ax=ax,
                    legend=False,
                    color='coral'
                )
                ax.set_xlabel('Nombre de valeurs manquantes')
                ax.set_ylabel('Variables')
                ax.set_title('Valeurs manquantes par variable')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close()
            except:
                pass
        else:
            st.success("Aucune valeur manquante détectée dans le dataset")
        
        # Tableau complet
        with st.expander("Voir toutes les variables"):
            st.dataframe(missing_df, use_container_width=True)


# Initialisation de la session state
if 'plan' not in st.session_state:
    st.session_state.plan = None
if 'plan_text' not in st.session_state:
    st.session_state.plan_text = ""
if 'csv_data' not in st.session_state:
    st.session_state.csv_data = None
if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False
if 'show_data_analysis' not in st.session_state:
    st.session_state.show_data_analysis = False

# Initialiser la langue par défaut
if 'language' not in st.session_state:
    st.session_state.language = 'fr'
if 'report_language' not in st.session_state:
    st.session_state.report_language = 'fr'

# Initialiser le mode d'analyse
if 'analysis_mode' not in st.session_state:
    st.session_state.analysis_mode = None

# Initialiser le contexte de l'étude
if 'study_context' not in st.session_state:
    st.session_state.study_context = None

# Initialiser le workflow
if 'workflow_step' not in st.session_state:
    st.session_state.workflow_step = 1
if 'workflow_history' not in st.session_state:
    st.session_state.workflow_history = [1]

# SÉLECTEUR DE LANGUE (en haut à droite)
col_title, col_lang = st.columns([5, 1])

with col_title:
    st.title("[DATA] AI Statistical Reporter")

with col_lang:
    lang = st.selectbox(
        "🌍",
        options=['fr', 'en'],
        format_func=get_language_name if TRANSLATIONS_AVAILABLE else lambda x: x,
        key='language',
        label_visibility="collapsed"
    )

st.markdown("---")

# Afficher le tour guidé pour les nouveaux utilisateurs
if UI_COMPONENTS_AVAILABLE:
    show_onboarding_tour()

# Proposer de récupérer un brouillon
if AUTOSAVE_AVAILABLE:
    show_draft_recovery()


# ========== SIDEBAR ==========
with st.sidebar:
    st.title("🎯 Statistical Reporter")
    st.markdown("---")
    
    st.markdown("### [DATA] Statut Configuration")
    
    # Vérifier si données chargées
    has_data = 'temp_path' in st.session_state and st.session_state.get('temp_path') and Path(st.session_state['temp_path']).exists()
    has_profile = st.session_state.get('writing_profile') is not None
    has_mode = st.session_state.get('use_context') is not None
    
    if has_data:
        filename = Path(st.session_state['temp_path']).name
        st.success(f"[OK] Données : {filename}")
    else:
        st.info("⚪ Aucune donnée")
    
    if has_profile:
        try:
            from study_context import WritingProfile
            profile = st.session_state.writing_profile
            if isinstance(profile, str):
                profile = WritingProfile(profile)
            profile_names = {
                WritingProfile.ACADEMIC: "Académique",
                WritingProfile.CONSULTANT: "Consultant",
                WritingProfile.INSTITUTIONAL: "Institutionnel"
            }
            st.success(f"[OK] Profil : {profile_names.get(profile, 'Sélectionné')}")
        except:
            st.success("[OK] Profil sélectionné")
    else:
        st.info("⚪ Aucun profil")
    
    if has_mode:
        mode = "Avec contexte" if st.session_state.use_context else "Sans contexte"
        st.success(f"[OK] Mode : {mode}")
    else:
        st.info("⚪ Aucun mode")
    
    st.markdown("---")
    st.info("📋 **Suivez les étapes sur la page principale** →")
    
    st.markdown("---")
    st.caption("© 2025 AI Statistical Reporter")





# ═══════════════════════════════════════════════════════════════════════════════
# WORKFLOW SÉQUENTIEL 5 ÉTAPES (PAGE PRINCIPALE)
# ═══════════════════════════════════════════════════════════════════════════════

# ========== WORKFLOW PRINCIPAL PAR ÉTAPES ==========

if not WORKFLOW_MANAGER_AVAILABLE:
    st.error("[ERROR] Module workflow_manager.py manquant")
    st.info("Téléchargez workflow_manager.py et placez-le dans le même dossier")
    st.stop()

current_step = st.session_state.get('workflow_step', 1)

# Barre de progression
step_names = {
    1: ("👤 Choix du profil", "Sélectionnez votre profil de rédaction"),
    2: ("🎯 Choix du mode", "Avez-vous une problématique définie ?"),
    2.5: ("📋 Contexte de l'étude", "Définissez votre contexte"),
    3: ("📂 Upload des données", "Chargez votre fichier CSV/Excel"),
    4: ("📝 Génération du plan", "Créez le plan de votre rapport"),
    5: ("📏 Configuration", "Définissez la longueur des chapitres"),
    6: ("📄 Génération du rapport", "Générez les chapitres")
}

if current_step == 2.5:
    progress = 2 / 6  # Étape 2.5 entre 2 et 3
else:
    progress = (current_step - 1) / 6  # Total 6 étapes principales

icon_title, description = step_names.get(current_step, ("Étape inconnue", ""))

st.markdown(f'''
<div style="margin-bottom: 2rem;">
    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem;">
        <span style="font-weight: 600; font-size: 1.2rem;">
            {icon_title}
        </span>
        <span style="color: #666; font-size: 0.9rem;">
            {int(progress * 100)}%
        </span>
    </div>
    <div style="background-color: #e0e0e0; border-radius: 10px; height: 10px; overflow: hidden;">
        <div style="background: linear-gradient(90deg, #4CAF50 0%, #45a049 100%); 
                    height: 100%; width: {progress * 100}%; transition: width 0.3s ease;">
        </div>
    </div>
    <p style="color: #666; font-size: 0.95rem; margin-top: 0.5rem; font-style: italic;">
        {description}
    </p>
</div>
''', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# ÉTAPE 1 : UPLOAD DES DONNÉES
# ═══════════════════════════════════════════════════════════════

if current_step == 1:
    st.title("👤 Étape 1/6 : Choix du profil de rédaction")
    
    st.info("💡 **Choisissez le style adapté à votre audience**")
    
    if STUDY_CONTEXT_AVAILABLE:
        try:
            from study_context import WritingProfile
            from writing_profiles import get_profile_summary
            
            summary = get_profile_summary()
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                profile = WritingProfile.ACADEMIC
                info = summary[profile]
                
                st.markdown(f"### {info['emoji']} {info['name']}")
                st.markdown(f'*"{info["phrase_cle"]}"*')
                st.markdown(f"**Pour :** {info['public']}")
                
                with st.expander("✨ Caractéristiques"):
                    for car in info['caracteristiques'][:4]:
                        st.markdown(f"• {car}")
                
                if st.button("Choisir Académique", type="primary", 
                           use_container_width=True, key="btn_academic"):
                    st.session_state.writing_profile = profile
                    st.session_state.workflow_step = 2
                    st.session_state.workflow_history.append(2)
                    st.rerun()
            
            with col2:
                profile = WritingProfile.CONSULTANT
                info = summary[profile]
                
                st.markdown(f"### {info['emoji']} {info['name']}")
                st.markdown(f'*"{info["phrase_cle"]}"*')
                st.markdown(f"**Pour :** {info['public']}")
                
                with st.expander("✨ Caractéristiques"):
                    for car in info['caracteristiques'][:4]:
                        st.markdown(f"• {car}")
                
                if st.button("Choisir Consultant", type="primary", 
                           use_container_width=True, key="btn_consultant"):
                    st.session_state.writing_profile = profile
                    st.session_state.workflow_step = 2
                    st.session_state.workflow_history.append(2)
                    st.rerun()
            
            with col3:
                profile = WritingProfile.INSTITUTIONAL
                info = summary[profile]
                
                st.markdown(f"### {info['emoji']} {info['name']}")
                st.markdown(f'*"{info["phrase_cle"]}"*')
                st.markdown(f"**Pour :** {info['public']}")
                
                with st.expander("✨ Caractéristiques"):
                    for car in info['caracteristiques'][:4]:
                        st.markdown(f"• {car}")
                
                if st.button("Choisir Institutionnel", type="primary", 
                           use_container_width=True, key="btn_institutional"):
                    st.session_state.writing_profile = profile
                    st.session_state.workflow_step = 2
                    st.session_state.workflow_history.append(2)
                    st.rerun()
        
        except Exception as e:
            st.error(f"[ERROR] Erreur chargement profils : {e}")
            if st.button("Continuer avec profil par défaut", type="secondary"):
                st.session_state.writing_profile = WritingProfile.ACADEMIC
                st.session_state.workflow_step = 2
                st.session_state.workflow_history.append(2)
                st.rerun()
    else:
        st.warning("[WARNING] Module profils non disponible")
        if st.button("Continuer avec profil par défaut", type="secondary"):
            st.session_state.writing_profile = "academic"
            st.session_state.workflow_step = 2
            st.session_state.workflow_history.append(2)
            st.rerun()

elif current_step == 2:
    st.title("🎯 Étape 2/6 : Choix du mode d'analyse")
    
    st.markdown('''
### Avez-vous un contexte d'étude défini ?

C'est-à-dire : une **problématique**, des **hypothèses**, des **objectifs** de recherche ?
    ''')
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('''
<div style="border: 2px solid #4CAF50; border-radius: 10px; padding: 1.5rem; min-height: 300px;">
    <h3 style="color: #4CAF50;">[OK] Oui, j'ai un contexte</h3>
    <p><strong>Mode Académique</strong></p>
    <ul>
        <li>Problématique définie</li>
        <li>Hypothèses à tester</li>
        <li>Objectifs clairs</li>
        <li>Analyse ciblée</li>
        <li>Rapport de haute qualité</li>
    </ul>
    <p style="color: #666; font-size: 0.9rem; margin-top: 1rem;">
        <em>📚 Recommandé pour : études sérieuses, mémoires, recherches, rapports professionnels</em>
    </p>
</div>
        ''', unsafe_allow_html=True)
        
        st.markdown("")
        if st.button("[OK] J'ai un contexte d'étude", type="primary", use_container_width=True, key="btn_academic"):
            st.session_state.analysis_mode = "academic"
            st.session_state.use_context = True
            if STUDY_CONTEXT_AVAILABLE:
                if st.session_state.study_context is None:
                    from study_context import StudyContext
                    st.session_state.study_context = StudyContext()
            st.session_state.workflow_step = 2.5
            st.session_state.workflow_history.append(2.5)
            st.rerun()
    
    with col2:
        st.markdown('''
<div style="border: 2px solid #2196F3; border-radius: 10px; padding: 1.5rem; min-height: 300px;">
    <h3 style="color: #2196F3;">⚡ Non, analyse rapide</h3>
    <p><strong>Mode Rapide</strong></p>
    <ul>
        <li>Pas de contexte prédéfini</li>
        <li>Analyse automatique</li>
        <li>Résultats génériques</li>
        <li>Rapide (5-10 minutes)</li>
        <li>Exploration des données</li>
    </ul>
    <p style="color: #666; font-size: 0.9rem; margin-top: 1rem;">
        <em>⚡ Recommandé pour : exploration rapide, prototypage, premiers insights</em>
    </p>
</div>
        ''', unsafe_allow_html=True)
        
        st.markdown("")
        if st.button("⚡ Analyse rapide sans contexte", use_container_width=True, key="btn_quick"):
            st.session_state.analysis_mode = "quick"
            st.session_state.study_context = None
            st.session_state.use_context =False
            st.session_state.workflow_step = 3
            st.session_state.workflow_history.append(3)
            st.rerun()
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← Retour", use_container_width=True):
            st.session_state.workflow_step = 1
            st.rerun()


# ═══════════════════════════════════════════════════════════════
# ÉTAPE 2.5 : CONTEXTE DE L'ÉTUDE (si mode académique)
# ═══════════════════════════════════════════════════════════════

elif current_step == 2.5:
    st.title("📋 Étape 2.5/6 : Contexte de l'étude")
    
    if not STUDY_CONTEXT_AVAILABLE:
        st.error("[ERROR] Module study_context.py non disponible")
        st.info("Téléchargez study_context.py et placez-le dans le même dossier")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            if st.button("← Retour", use_container_width=True):
                st.session_state.workflow_step = 2
                st.rerun()
        with col3:
            if st.button("Passer cette étape →", use_container_width=True):
                st.session_state.workflow_step = 3
                st.session_state.workflow_history.append(3)
                st.rerun()
        st.stop()
    
    from study_context import StudyContext
    
    st.info('''
**Définissez votre contexte pour une analyse de qualité académique**

Remplissez au minimum la question de recherche OU les objectifs.
    ''')
    
    if st.session_state.study_context is None:
        st.session_state.study_context = StudyContext()
    
    ctx = st.session_state.study_context
    
    ctx.study_title = st.text_input(
        "📌 Titre de l'étude",
        value=ctx.study_title,
        placeholder="Ex: Analyse de la satisfaction client 2024"
    )
    
    ctx.research_question = st.text_area(
        "❓ Question de recherche principale ⭐",
        value=ctx.research_question,
        height=100,
        placeholder="Ex: Quel est l'impact de l'âge sur le salaire ?"
    )
    
    st.markdown("**🎯 Objectifs** *(un par ligne)* ⭐")
    obj_text = st.text_area(
        "Objectifs",
        value="\n".join(ctx.objectives) if ctx.objectives else "",
        height=100,
        placeholder="1. Analyser\n2. Identifier\n3. Proposer",
        label_visibility="collapsed"
    )
    ctx.objectives = [o.strip() for o in obj_text.split('\n') if o.strip()]
    
    is_valid = bool(ctx.research_question or ctx.objectives)
    
    if ctx.research_question and ctx.objectives:
        st.success("[OK] Contexte minimal rempli !")
    elif is_valid:
        st.info("[INFO] Vous pouvez continuer")
    else:
        st.warning("[WARNING] Remplissez au moins la question ou les objectifs")
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← Retour", use_container_width=True, key="back_context"):
            st.session_state.workflow_step = 2
            st.rerun()
    with col3:
        if st.button(
            "Valider et continuer →",
            type="primary",
            use_container_width=True,
            disabled=not is_valid,
            key="next_context"
        ):
            st.session_state.workflow_step = 3
            st.session_state.workflow_history.append(3)
            st.rerun()

# ═══════════════════════════════════════════════════════════════
# ÉTAPE 3 : GÉNÉRATION DU PLAN - VERSION COMPLÈTE AVEC 3 OPTIONS
# ═══════════════════════════════════════════════════════════════
elif current_step == 3:
    st.title("📂 Étape 3/6 : Upload des données")
    
    st.info("""
💡 **Chargez maintenant vos données statistiques**

Uploadez votre fichier CSV ou Excel ci-dessous.
    """)
    
    # ═══════════════════════════════════════════════════════════════
    # WIDGET UPLOAD DIRECTEMENT DANS L'ÉTAPE 3
    # ═══════════════════════════════════════════════════════════════
    
    uploaded_file = st.file_uploader(
        "📤 Glissez-déposez ou cliquez pour parcourir",
        type=['csv', 'xlsx', 'xls'],
        help="Formats acceptés : CSV, Excel (.xlsx, .xls)",
        key="step3_uploader"
    )
    
    # Si fichier uploadé, le traiter
    if uploaded_file is not None:
        try:
            # Sauvegarder temporairement
            temp_path = Path("temp") / uploaded_file.name
            temp_path.parent.mkdir(exist_ok=True)
            
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Charger selon le type
            file_ext = temp_path.suffix.lower()
            
            if file_ext == '.csv':
                try:
                    df = pd.read_csv(temp_path, encoding='utf-8')
                except UnicodeDecodeError:
                    try:
                        df = pd.read_csv(temp_path, encoding='latin-1')
                    except:
                        df = pd.read_csv(temp_path, encoding='iso-8859-1')
            
            elif file_ext in ['.xlsx', '.xls']:
                try:
                    df = pd.read_excel(temp_path, engine='openpyxl' if file_ext == '.xlsx' else None)
                except ImportError:
                    st.error("[ERROR] Installation requise : `pip install openpyxl`")
                    st.stop()
                except Exception as e:
                    st.error(f"[ERROR] Erreur lors de la lecture Excel : {e}")
                    st.stop()
            
            # Stocker dans session_state
            st.session_state.csv_data = df
            st.session_state['temp_path'] = str(temp_path)
            st.session_state['uploaded_filename'] = uploaded_file.name
            
            st.success(f"[OK] Fichier **{uploaded_file.name}** chargé avec succès !")
            
        except Exception as e:
            st.error(f"[ERROR] Erreur lors du chargement : {e}")
            st.stop()
    
    # ═══════════════════════════════════════════════════════════════
    # AFFICHAGE DES DONNÉES SI DÉJÀ CHARGÉES
    # ═══════════════════════════════════════════════════════════════
    
    if st.session_state.csv_data is not None:
        df = st.session_state.csv_data
        
        st.markdown("---")
        
        # Métriques
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("[DATA] Lignes", f"{len(df):,}")
        with col2:
            st.metric("📋 Colonnes", len(df.columns))
        with col3:
            size_mb = df.memory_usage(deep=True).sum() / 1024 / 1024
            st.metric("[SAVE] Taille", f"{size_mb:.2f} MB")
        
        # Aperçu
        with st.expander("[SEARCH] Aperçu des données", expanded=True):
            st.dataframe(df.head(10), use_container_width=True)
        
        # Statistiques
        with st.expander("[DATA] Statistiques descriptives"):
            st.dataframe(df.describe().T, use_container_width=True)
        
        # Informations sur les colonnes
        with st.expander("[INFO] Informations sur les colonnes"):
            col_info = pd.DataFrame({
                'Type': df.dtypes,
                'Valeurs manquantes': df.isnull().sum(),
                '% manquant': (df.isnull().sum() / len(df) * 100).round(2)
            })
            st.dataframe(col_info, use_container_width=True)
        
        st.markdown("---")
        
        # Navigation
        col1, col2, col3 = st.columns([1, 2, 1])
        with col3:
            if st.button("➡️ Suivant : Génération du plan", type="primary", use_container_width=True):
                st.session_state.workflow_step = 4
                st.session_state.workflow_history.append(4)
                st.rerun()
    
    else:
        # Message si aucun fichier
        st.markdown("---")
        st.info("""
### 📋 Instructions
        
1. **Cliquez** sur le bouton ci-dessus ou **glissez-déposez** votre fichier
2. Les formats acceptés sont : **CSV** (.csv), **Excel** (.xlsx, .xls)
3. Le fichier sera automatiquement analysé
4. Vous verrez un aperçu des données avant de continuer

**Conseil** : Assurez-vous que votre fichier contient des en-têtes de colonnes.
        """)
elif current_step == 4:
    st.title("📝 Étape 4/6 : Génération du plan")

    # ═══════════════════════════════════════════════════════════════
    # VÉRIFICATION DES PRÉREQUIS
    # ═══════════════════════════════════════════════════════════════
    
    if st.session_state.csv_data is None:
        st.error("[ERROR] Aucune donnée chargée")
        st.info("👈 Retournez à l'étape 3 pour uploader vos données")
        
        if st.button("← Retour à l'étape 3", type="primary"):
            st.session_state.workflow_step = 3
            st.rerun()
        st.stop()
    
    # ═══════════════════════════════════════════════════════════════
    # AFFICHAGE CONTEXTE : PROFIL + MODE + DONNÉES
    # ═══════════════════════════════════════════════════════════════
    
    st.markdown("### 🎯 Configuration actuelle")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Afficher profil
        profile = st.session_state.get('writing_profile')
        if profile:
            if STUDY_CONTEXT_AVAILABLE:
                try:
                    from study_context import WritingProfile
                    from writing_profiles import get_profile_summary
                    
                    # Convertir en WritingProfile si c'est une string
                    if isinstance(profile, str):
                        profile = WritingProfile(profile)
                    
                    summary = get_profile_summary()
                    info = summary.get(profile, {})
                    
                    st.info(f"**👤 Profil**\n\n{info.get('emoji', '📝')} **{info.get('name', profile)}**")
                except Exception as e:
                    st.info(f"**👤 Profil**\n\n{profile}")
            else:
                st.info(f"**👤 Profil**\n\n{profile}")
        else:
            st.warning("[WARNING] **Profil non défini**")
    
    with col2:
        # Afficher mode
        use_context = st.session_state.get('use_context')
        if use_context is not None:
            if use_context:
                st.info("**📚 Mode**\n\n📝 Avec contexte")
            else:
                st.info("**📚 Mode**\n\n⚡ Sans contexte")
        else:
            st.warning("[WARNING] **Mode non défini**")
    
    with col3:
        # Afficher données
        df = st.session_state.csv_data
        st.info(f"**[DATA] Données**\n\n{len(df):,} lignes\n{len(df.columns)} colonnes")
    
    st.markdown("---")
    
    # ═══════════════════════════════════════════════════════════════
    # MESSAGE CONTEXTUEL SELON LE PROFIL
    # ═══════════════════════════════════════════════════════════════
    
    profile = st.session_state.get('writing_profile')
    
    if profile and STUDY_CONTEXT_AVAILABLE:
        try:
            from study_context import WritingProfile
            from writing_profiles import get_profile_summary
            
            # Convertir en WritingProfile si nécessaire
            if isinstance(profile, str):
                profile = WritingProfile(profile)
            
            summary = get_profile_summary()
            info = summary.get(profile, {})
            
            # Messages adaptés par profil
            if profile == WritingProfile.ACADEMIC:
                st.info("""
💡 **Plan adapté au profil Académique**

Le plan généré sera structuré pour :
- [OK] Répondre à une **question de recherche**
- [OK] Tester des **hypothèses**
- [OK] Respecter les **standards académiques**
- [OK] Inclure une **revue théorique**
- [OK] Présenter une **méthodologie rigoureuse**
                """)
            
            elif profile == WritingProfile.CONSULTANT:
                st.info("""
💡 **Plan adapté au profil Consultant**

Le plan généré sera structuré pour :
- [OK] Identifier les **opportunités business**
- [OK] Fournir des **recommandations actionnables**
- [OK] Mettre l'accent sur les **insights clés**
- [OK] Quantifier les **impacts potentiels**
- [OK] Proposer un **plan d'action concret**

**Inclura automatiquement** :
- 📋 Executive Summary
- 🎯 Chapitre dédié aux Recommandations Stratégiques
                """)
            
            elif profile == WritingProfile.INSTITUTIONAL:
                st.info("""
💡 **Plan adapté au profil Institutionnel**

Le plan généré sera structuré pour :
- [OK] Garantir la **transparence** et la **traçabilité**
- [OK] Respecter les **obligations réglementaires**
- [OK] Documenter les **processus** en détail
- [OK] Faciliter la **reproductibilité**
- [OK] Maintenir un ton **formel et neutre**
                """)
        except Exception as e:
            st.info("💡 Le plan sera généré en fonction de vos données et de votre profil sélectionné")
    else:
        st.info("💡 Génération du plan du rapport basé sur vos données" +
                (" et votre contexte" if st.session_state.get('use_context') else ""))

    # ═══════════════════════════════════════════════════════════════
    # BOUTON DE GÉNÉRATION
    # ═══════════════════════════════════════════════════════════════
    
    col1, col2 = st.columns([3, 1])

    with col1:
        st.markdown("### 🚀 Générer le plan")
        profile_name = "standard"
        if profile and STUDY_CONTEXT_AVAILABLE:
            try:
                from study_context import WritingProfile
                if isinstance(profile, WritingProfile):
                    profile_name = profile.value
                elif isinstance(profile, str):
                    profile_name = profile
                
                st.caption(f"L'assistant va créer un plan **adapté au profil {profile_name.capitalize()}**")
            except:
                st.caption("L'assistant va analyser vos données et proposer une structure de rapport")
        else:
            st.caption("L'assistant va analyser vos données et proposer une structure de rapport")

    with col2:
        generate_btn = st.button(
            "🚀 Générer",
            type="primary",
            use_container_width=True,
            key="generate_plan_btn"
        )

    # ═══════════════════════════════════════════════════════════════
    # GÉNÉRATION DU PLAN AVEC PROFIL
    # ═══════════════════════════════════════════════════════════════
    
    if generate_btn:
        with st.spinner("[REFRESH] Génération du plan en cours... (30-60 secondes)"):
            try:
                from week2_architect_agent import analyze_csv, generate_report_plan

                # Récupérer le chemin du fichier
                csv_path = st.session_state.get('temp_path')
                
                if not csv_path:
                    st.error("[ERROR] Chemin du fichier non trouvé")
                    st.stop()
                
                # Analyser les métadonnées
                with st.spinner("[DATA] Analyse des métadonnées..."):
                    metadata = analyze_csv(csv_path)
                    st.success(f"[OK] Analyse terminée : {metadata['shape']['rows']:,} lignes analysées")

                # Récupérer le profil de rédaction
                writing_profile = st.session_state.get('writing_profile')
                
                # Convertir en string si c'est un WritingProfile
                if writing_profile and STUDY_CONTEXT_AVAILABLE:
                    try:
                        from study_context import WritingProfile
                        if isinstance(writing_profile, WritingProfile):
                            writing_profile = writing_profile.value
                    except:
                        pass
                
                # Récupérer le contexte d'étude si disponible
                study_ctx = st.session_state.get('study_context', None) if STUDY_CONTEXT_AVAILABLE else None

                # ═══════════════════════════════════════════════════════════
                # GÉNÉRER LE PLAN AVEC PROFIL ⭐
                # ═══════════════════════════════════════════════════════════
                
                with st.spinner(f"✍️ Génération du plan {writing_profile or 'standard'}..."):
                    try:
                        # Essayer avec profil ET contexte
                        plan = generate_report_plan(
                            metadata, 
                            study_context=study_ctx,
                            writing_profile=writing_profile
                        )
                    except TypeError:
                        # Fallback : essayer sans writing_profile
                        try:
                            plan = generate_report_plan(metadata, study_context=study_ctx)
                            st.warning("[WARNING] Le profil de rédaction n'a pas pu être pris en compte. Veuillez mettre à jour `week2_architect_agent.py`")
                        except TypeError:
                            # Fallback final : juste metadata
                            plan = generate_report_plan(metadata)
                            st.warning("[WARNING] Ni le profil ni le contexte n'ont pu être pris en compte")

                # Sauvegarder dans session_state
                st.session_state.plan = plan
                st.session_state.plan_text = json_to_editable_text(plan)
                
                # Afficher succès avec détails
                st.success("[OK] Plan généré avec succès !")
                
                # Afficher statistiques du plan
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("📚 Chapitres", len(plan.get('chapitres', [])))
                with col2:
                    total_sections = sum(len(chap.get('sections', [])) for chap in plan.get('chapitres', []))
                    st.metric("📑 Sections", total_sections)
                with col3:
                    profil_plan = plan.get('profil', 'N/A')
                    st.metric("👤 Profil", profil_plan.capitalize() if profil_plan != 'N/A' else 'Standard')
                
                # Si profil consultant, mentionner les chapitres clés
                if writing_profile and writing_profile.lower() == "consultant":
                    st.info("💼 **Plan Consultant généré** : Inclut Executive Summary et Recommandations Stratégiques")
                
                st.rerun()

            except Exception as e:
                st.error(f"[ERROR] Erreur lors de la génération du plan : {str(e)}")
                
                # Afficher détails de l'erreur
                with st.expander("[SEARCH] Détails de l'erreur"):
                    st.exception(e)
                
                # Suggestions
                st.info("""
💡 **Solutions possibles** :
1. Vérifiez que la clé API `GMINI_API_KEY` est configurée
2. Vérifiez que le fichier `week2_architect_agent.py` est à jour
3. Vérifiez votre connexion internet
4. Essayez de régénérer le plan
                """)

    # ════════════════════════════════════════════════════════════
    # AFFICHAGE ET MODIFICATION DU PLAN (3 OPTIONS)
    # ════════════════════════════════════════════════════════════
    
    if st.session_state.get('plan'):
        st.success("[OK] Plan disponible")
        
        # Initialiser le mode si pas défini
        if 'plan_action_mode' not in st.session_state:
            st.session_state.plan_action_mode = 'view'
        
        # ────────────────────────────────────────────────────────
        # BOUTONS D'ACTION (3 OPTIONS)
        # ────────────────────────────────────────────────────────
        
        st.markdown("### 📋 Votre plan de rapport")
        
        col_view, col_edit, col_regen = st.columns(3)
        
        with col_view:
            if st.button("👁️ Voir le plan", use_container_width=True, key="btn_view_plan", 
                       type="primary" if st.session_state.plan_action_mode == 'view' else "secondary"):
                st.session_state.plan_action_mode = 'view'
                st.rerun()
        
        with col_edit:
            if st.button("✏️ Éditer manuellement", use_container_width=True, key="btn_edit_plan",
                       type="primary" if st.session_state.plan_action_mode == 'edit' else "secondary"):
                st.session_state.plan_action_mode = 'edit'
                if 'plan_text' not in st.session_state or not st.session_state.plan_text:
                    st.session_state.plan_text = json_to_editable_text(st.session_state.plan)
                st.rerun()
        
        with col_regen:
            if st.button("[REFRESH] Régénérer avec IA", use_container_width=True, key="btn_regen_plan",
                       type="primary" if st.session_state.plan_action_mode == 'regenerate' else "secondary"):
                st.session_state.plan_action_mode = 'regenerate'
                st.rerun()
        
        st.markdown("---")
        
        # ════════════════════════════════════════════════════════
        # OPTION 1 : VISUALISATION SIMPLE
        # ════════════════════════════════════════════════════════
        
        if st.session_state.plan_action_mode == 'view':
            st.info("👁️ **Mode visualisation** - Consultez votre plan")
            
            with st.expander("📋 Plan détaillé", expanded=True):
                plan = st.session_state.plan

                if isinstance(plan, dict):
                    st.markdown(f"### {plan.get('titre', 'Plan du Rapport')}")
                    st.caption(f"📅 Date: {plan.get('date', '')} | ✍️ Auteur: {plan.get('auteur', '')}")
                    
                    st.markdown("---")
                    
                    # Statistiques du plan
                    total_sections = sum(len(chap.get('sections', [])) for chap in plan.get('chapitres', []))
                    total_analyses = sum(
                        len(sec.get('analyses', [])) 
                        for chap in plan.get('chapitres', []) 
                        for sec in chap.get('sections', [])
                    )
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("📚 Chapitres", len(plan.get('chapitres', [])))
                    with col2:
                        st.metric("📑 Sections", total_sections)
                    with col3:
                        st.metric("[SEARCH] Analyses", total_analyses)
                    
                    st.markdown("---")

                    for i, chap in enumerate(plan.get('chapitres', []), 1):
                        with st.expander(f"**{i}. {chap.get('titre')}**", expanded=False):
                            for j, sec in enumerate(chap.get('sections', []), 1):
                                st.markdown(f"**{i}.{j}. {sec.get('titre')}**")
                                
                                analyses = sec.get('analyses', [])
                                for analyse in analyses:
                                    st.markdown(f"   • {analyse}")
                                
                                if j < len(chap.get('sections', [])):
                                    st.markdown("")
                else:
                    st.json(plan)
        
        # ════════════════════════════════════════════════════════
        # OPTION 2 : ÉDITION MANUELLE
        # ════════════════════════════════════════════════════════
        
        elif st.session_state.plan_action_mode == 'edit':
            st.info("✏️ **Mode édition manuelle** - Modifiez le texte directement")
            
            st.markdown("""
**Instructions :**
- Gardez le format : `TITRE:`, `DATE:`, `AUTEUR:`
- Chapitres : `1. Titre du chapitre`
- Sections : `   1.1. Titre de la section` (indenté)
- Analyses : `      - Analyse à faire` (indenté)
            """)
            
            # Zone de texte pour éditer
            if 'plan_text' not in st.session_state or not st.session_state.plan_text:
                st.session_state.plan_text = json_to_editable_text(st.session_state.plan)
            
            edited_text = st.text_area(
                "Éditez le plan :",
                value=st.session_state.plan_text,
                height=400,
                key="plan_editor"
            )
            
            st.markdown("---")
            col_save, col_cancel = st.columns(2)
            
            with col_save:
                if st.button("[SAVE] Enregistrer les modifications", type="primary", use_container_width=True):
                    try:
                        with st.spinner("Conversion du texte en plan structuré..."):
                            # Parser le texte avec l'IA
                            new_plan = text_to_json_with_ai(edited_text)
                            
                            # Mettre à jour le plan
                            st.session_state.plan = new_plan
                            st.session_state.plan_text = edited_text
                            st.session_state.plan_action_mode = 'view'
                            
                            st.success("[OK] Plan mis à jour avec succès !")
                            
                            if LOGGING_AVAILABLE:
                                log_user_action('plan_edited_manually', {
                                    'chapters': len(new_plan.get('chapitres', []))
                                })
                            
                            st.rerun()
                    
                    except Exception as e:
                        st.error(f"[ERROR] Erreur : {str(e)}")
                        st.info("💡 Vérifiez le format du texte")
            
            with col_cancel:
                if st.button("[ERROR] Annuler", use_container_width=True):
                    st.session_state.plan_action_mode = 'view'
                    st.rerun()
        
        # ════════════════════════════════════════════════════════
        # OPTION 3 : RÉGÉNÉRATION AVEC IA
        # ════════════════════════════════════════════════════════
        
        elif st.session_state.plan_action_mode == 'regenerate':
            st.info("[REFRESH] **Mode régénération intelligente** - L'IA va améliorer votre plan selon vos instructions")
            
            st.markdown("""
**💡 Exemples d'instructions :**
- "Ajouter un chapitre sur l'analyse des outliers"
- "Approfondir la section sur les corrélations"
- "Supprimer le chapitre 3 et fusionner son contenu avec le chapitre 2"
- "Simplifier le vocabulaire pour un public non-technique"
            """)
            
            # Zone de texte pour les instructions
            regen_instructions = st.text_area(
                "📝 Décrivez les améliorations souhaitées :",
                height=150,
                placeholder="Ex: Ajouter un chapitre sur les tests statistiques...",
                key="regen_instructions"
            )
            
            # Options avancées
            st.markdown("**⚙️ Options avancées :**")
            
            col_opt1, col_opt2, col_opt3 = st.columns(3)
            
            with col_opt1:
                keep_structure = st.checkbox(
                    "Conserver la structure actuelle",
                    value=False,
                    help="Garde le même nombre de chapitres et leur ordre"
                )
            
            with col_opt2:
                academic_mode = st.checkbox(
                    "Style académique",
                    value=st.session_state.get('analysis_mode') == "academic",
                    help="Adapte le plan pour un rapport académique"
                )
            
            with col_opt3:
                detailed_mode = st.checkbox(
                    "Mode détaillé",
                    value=False,
                    help="Génère plus de sections et d'analyses par chapitre"
                )
            
            st.markdown("---")
            
            col_regen, col_cancel = st.columns(2)
            
            with col_regen:
                if st.button("🚀 Régénérer le plan", type="primary", use_container_width=True, 
                           disabled=not regen_instructions.strip()):
                    
                    with st.spinner("Régénération du plan en cours... (30-60 secondes)"):
                        try:
                            from week2_architect_agent import analyze_csv
                            
                            # [OK] CORRECTION : Utiliser temp_path
                            csv_path = st.session_state.get('temp_path')
                            
                            if not csv_path:
                                st.error("[ERROR] Chemin du fichier non trouvé")
                                st.stop()
                            
                            # Récupérer les métadonnées
                            metadata = analyze_csv(csv_path)
                            
                            # Construire le prompt de régénération
                            current_plan_text = json_to_editable_text(st.session_state.plan)
                            
                            # Ajouter le contexte d'étude si disponible
                            study_ctx = st.session_state.get('study_context', None) if STUDY_CONTEXT_AVAILABLE else None
                            
                            # Régénérer le plan
                            try:
                                new_plan = regenerate_plan_with_instructions(
                                    current_plan=st.session_state.plan,
                                    instructions=regen_instructions,
                                    metadata=metadata,
                                    keep_structure=keep_structure,
                                    academic=academic_mode,
                                    detailed=detailed_mode,
                                    study_context=study_ctx
                                )
                            except TypeError:
                                new_plan = regenerate_plan_with_instructions(
                                    current_plan=st.session_state.plan,
                                    instructions=regen_instructions,
                                    metadata=metadata,
                                    keep_structure=keep_structure,
                                    academic=academic_mode,
                                    detailed=detailed_mode,
                                    study_context=None
                                )
                            
                            # Mettre à jour
                            st.session_state.plan = new_plan
                            st.session_state.plan_text = json_to_editable_text(new_plan)
                            st.session_state.plan_action_mode = 'view'
                            
                            st.success("[OK] Plan régénéré avec succès !")
                            
                            if LOGGING_AVAILABLE:
                                log_user_action('plan_regenerated', {
                                    'instructions': regen_instructions[:100],
                                    'keep_structure': keep_structure,
                                    'academic': academic_mode,
                                    'detailed': detailed_mode
                                })
                            
                            st.rerun()
                        
                        except Exception as e:
                            st.error(f"[ERROR] Erreur lors de la régénération : {str(e)}")
                            st.exception(e)
            
            with col_cancel:
                if st.button("[ERROR] Annuler", use_container_width=True):
                    st.session_state.plan_action_mode = 'view'
                    st.rerun()
    
    # ════════════════════════════════════════════════════════
    # NAVIGATION
    # ════════════════════════════════════════════════════════
    
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.button("← Retour", use_container_width=True, key="back_plan"):
            st.session_state.workflow_step = 3  # Retour vers Upload
            st.rerun()
    
    with col3:
        plan_ready = st.session_state.get('plan') is not None
        if st.button(
            "Continuer vers la configuration →",
            type="primary",
            use_container_width=True,
            disabled=not plan_ready,
            key="next_plan"
        ):
            st.session_state.workflow_step = 5  # [OK] CORRECTION : 5 au lieu de 4
            st.session_state.workflow_history.append(5)
            st.rerun()
# ═══════════════════════════════════════════════════════════════
# ÉTAPE 4 : CONFIGURATION
# ═══════════════════════════════════════════════════════════════

elif current_step == 5:
    st.title("📏 Étape 5/6 : Configuration du rapport")
    
    if st.session_state.plan is None:
        st.warning("[WARNING] Veuillez d'abord générer un plan")
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            if st.button("← Retour au plan", use_container_width=True):
                st.session_state.workflow_step = 3
                st.rerun()
        st.stop()
    
    st.info("""
**Configurez la longueur de chaque chapitre**

Définissez le nombre de pages souhaité pour chaque chapitre (1-30 pages).
Chaque page contient environ 300 mots.
    """)
    
    if COST_CONTROLLER_AVAILABLE:
        from cost_controller import cost_controller, display_cost_summary_in_streamlit
        
        if 'cost_controller' not in st.session_state:
            st.session_state.cost_controller = cost_controller
        
        cost_ctrl = st.session_state.cost_controller
        
        st.markdown("### 📏 Longueur des chapitres")
        
        plan = st.session_state.plan
        
        if isinstance(plan, dict) and 'chapitres' in plan:
            chapitres = plan['chapitres']
            
            for i, chapitre in enumerate(chapitres, 1):
                chapter_title = chapitre.get('titre', f'Chapitre {i}')
                
                with st.expander(f"📖 Chapitre {i} : {chapter_title}", expanded=(i==1)):
                    pages = st.slider(
                        f"Nombre de pages",
                        min_value=1,
                        max_value=30,
                        value=5,
                        key=f"pages_ch{i}_config"
                    )
                    
                    words = pages * 300
                    st.caption(f"📝 Environ {words:,} mots")
                    
                    cost_ctrl.set_chapter_length(i, chapter_title, pages)
        
        st.markdown("---")
        st.markdown("### 💰 Estimation des coûts")
        
        try:
            display_cost_summary_in_streamlit(cost_ctrl)
        except Exception as e:
            st.info("Configuration enregistrée")
        
        st.session_state.cost_controller = cost_ctrl
        st.session_state.chapter_lengths_configured = True
    
    else:
        st.warning("[WARNING] Module cost_controller.py non disponible")
        st.info("Les chapitres seront générés avec une longueur par défaut (5 pages par chapitre)")
        st.session_state.chapter_lengths_configured = True
    
    # Navigation
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.button("← Retour au plan", use_container_width=True, key="back_config"):
            st.session_state.workflow_step = 3
            st.rerun()
    
    with col3:
        if st.button(
            "🚀 Lancer la génération →",
            type="primary",
            use_container_width=True,
            key="next_config"
        ):
            #current_step==
            st.session_state.workflow_step = 6
            st.session_state.workflow_history.append(6)
            st.rerun()


# ═══════════════════════════════════════════════════════════════
# ÉTAPE 6 : GÉNÉRATION DU RAPPORT
# ═══════════════════════════════════════════════════════════════

elif current_step == 6:
    st.title("📄 Étape 6/6 : Génération du rapport")
    st.markdown("*Génération chapitre par chapitre avec validation humaine*")
    
    if st.session_state.plan is None:
        st.warning("[WARNING] Vous devez d'abord générer un plan dans la page 'Génération du plan'")
    if st.session_state.plan is None:
        st.warning("[WARNING] Vous devez d'abord générer un plan à l'étape 4")
        
        if st.button("← Retour à l'étape 4", type="primary"):
            st.session_state.workflow_step = 4
            st.rerun()
        st.stop()
    
    if not WORKFLOW_AVAILABLE:
        st.error("[ERROR] Le module Chapter Workflow n'est pas disponible")
        st.info("Installez les dépendances : `pip install e2b google-generativeai`")
        st.stop()
    
    # [OK] CORRECTION : Vérifier que temp_path existe
    if not st.session_state.get('temp_path'):
        st.error("[ERROR] Chemin du fichier CSV non trouvé")
        st.info("Retournez à l'étape 3 pour uploader vos données")
        
        if st.button("← Retour à l'étape 3", type="primary"):
            st.session_state.workflow_step = 3
            st.rerun()
        st.stop()
    
    # ═══════════════════════════════════════════════════════════════
    # AFFICHAGE CONTEXTE
    # ═══════════════════════════════════════════════════════════════
    
    st.markdown("### 🎯 Configuration actuelle")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Profil
        profile = st.session_state.get('writing_profile')
        if profile:
            if STUDY_CONTEXT_AVAILABLE:
                try:
                    from study_context import WritingProfile
                    from writing_profiles import get_profile_summary
                    
                    if isinstance(profile, str):
                        profile = WritingProfile(profile)
                    
                    summary = get_profile_summary()
                    info = summary.get(profile, {})
                    
                    st.info(f"**👤 Profil**\n\n{info.get('emoji', '📝')} {info.get('name', profile)}")
                except:
                    st.info(f"**👤 Profil**\n\n{profile}")
            else:
                st.info(f"**👤 Profil**\n\n{profile}")
        else:
            st.warning("[WARNING] Profil non défini")
    
    with col2:
        # Mode
        use_context = st.session_state.get('use_context')
        if use_context is not None:
            if use_context:
                st.info("**📚 Mode**\n\n📝 Avec contexte")
            else:
                st.info("**📚 Mode**\n\n⚡ Sans contexte")
        else:
            st.warning("[WARNING] Mode non défini")
    
    with col3:
        # Données
        df = st.session_state.csv_data
        st.info(f"**[DATA] Données**\n\n{len(df):,} lignes\n{len(df.columns)} colonnes")
    
    st.markdown("---")
    #elif not WORKFLOW_AVAILABLE:
        #st.error("[ERROR] Le module Chapter Workflow n'est pas disponible")
        #st.info("Installez les dépendances : `pip install e2b google-generativeai`")
    
    #else:
        # === NOUVELLE SECTION : CONFIGURATION DE LA LONGUEUR DES CHAPITRES ===
    if 'chapter_lengths_configured' not in st.session_state:
            st.session_state.chapter_lengths_configured = False
        
    if not st.session_state.chapter_lengths_configured:
            st.markdown("---")
            st.markdown("### 📏 Configuration de la longueur des chapitres")
            
            st.info("""
💡 **Contrôlez les coûts et le niveau de détail**

Définissez le nombre de pages souhaité pour chaque chapitre (1-30 pages).
            
**Guide de longueur :**
- **1-3 pages** : Synthèse concise (~300-900 mots) - Idéal pour intro/conclusion
- **4-7 pages** : Analyse standard (~1,200-2,100 mots) - Bon équilibre
- **8-15 pages** : Analyse détaillée (~2,400-4,500 mots) - Analyses approfondies
- **16-30 pages** : Analyse exhaustive (~4,800-9,000 mots) - Très détaillé

**Note :** 1 page ≈ 300 mots (standard académique)
            """)
            
            # Configuration par chapitre
            if COST_CONTROLLER_AVAILABLE:
                from cost_controller import cost_controller
                
                st.markdown("#### [DATA] Longueur par chapitre")
                
                # Récupérer les chapitres du plan
                plan = st.session_state.plan
                chapters_data = plan.get('chapitres', [])
                
                # Créer une interface pour chaque chapitre
                for chapter_data in chapters_data:
                    chapter_num = chapter_data['numero']
                    chapter_title = chapter_data['titre']
                    
                    col1, col2, col3 = st.columns([4, 2, 2])
                    
                    with col1:
                        st.markdown(f"**Chapitre {chapter_num}** : {chapter_title}")
                    
                    with col2:
                        # Déterminer la valeur par défaut selon le type de chapitre
                        default_pages = 5  # Par défaut
                        title_lower = chapter_title.lower()
                        
                        if any(word in title_lower for word in ['introduction', 'présentation', 'contexte']):
                            default_pages = 3
                        elif any(word in title_lower for word in ['conclusion', 'synthèse', 'recommandation']):
                            default_pages = 2
                        elif any(word in title_lower for word in ['analyse', 'étude', 'exploration']):
                            default_pages = 8
                        elif any(word in title_lower for word in ['modélisation', 'prédiction', 'test']):
                            default_pages = 10
                        
                        # Slider pour définir les pages
                        pages = st.slider(
                            "Pages",
                            min_value=1,
                            max_value=30,
                            value=default_pages,
                            step=1,
                            key=f"pages_ch{chapter_num}",
                            help=f"Définir la longueur du chapitre {chapter_num}"
                        )
                    
                    with col3:
                        # Afficher les mots estimés
                        words = pages * 300
                        st.caption(f"📝 ~{words:,} mots")
                    
                    # Sauvegarder dans le cost controller
                    cost_controller.set_chapter_length(chapter_num, chapter_title, pages)
                
                # Afficher le résumé des coûts
                st.markdown("---")
                from cost_controller import display_cost_summary_in_streamlit
                display_cost_summary_in_streamlit(cost_controller)
                
                # Sauvegarder dans session state
                st.session_state.cost_controller = cost_controller
                
                # Bouton pour commencer la génération
                st.markdown("---")
                if st.button("🚀 Commencer la génération avec ces paramètres", type="primary", use_container_width=True):
                    st.session_state.chapter_lengths_configured = True
                    st.rerun()
            
            else:
                # Si cost_controller pas disponible, bouton simple
                st.warning("[WARNING] Module cost_controller.py non disponible - Longueur par défaut utilisée")
                if st.button("🚀 Commencer la génération", type="primary", use_container_width=True):
                    st.session_state.chapter_lengths_configured = True
                    st.rerun()
        
        # === FIN NOUVELLE SECTION ===
        
        # Initialiser le workflow si nécessaire (UNIQUEMENT après configuration)
    if st.session_state.chapter_lengths_configured and 'workflow' not in st.session_state:
            if WORKFLOW_AVAILABLE:
                try:
                    user_id = st.session_state.get('user_id', 'default')
                    csv_path = st.session_state.get('temp_path')
                    # Passer le cost_controller et le study_context au workflow
                    cost_ctrl = st.session_state.get('cost_controller', None) if COST_CONTROLLER_AVAILABLE else None
                    study_ctx = st.session_state.get('study_context', None) if STUDY_CONTEXT_AVAILABLE else None
                    st.session_state.workflow = initialize_workflow(
                        user_id=user_id,
                        plan=st.session_state.plan,
                        csv_path=csv_path,
                        cost_controller=cost_ctrl,
                        study_context=study_ctx  # NOUVEAU paramètre
                    )
                    
                    if LOGGING_AVAILABLE:
                        logger.info(f"Workflow initialized: {len(st.session_state.workflow.chapters)} chapters")
                
                except Exception as e:
                    st.error(f"[ERROR] Erreur initialisation workflow: {e}")
                    st.stop()
        
        # Vérifier que le workflow est bien initialisé
    if 'workflow' not in st.session_state:
            st.info("[WARNING] Veuillez configurer la longueur des chapitres ci-dessus pour commencer.")
            st.stop()
        
    workflow = st.session_state.workflow
        
        # Afficher la progression
    st.markdown("### [DATA] Progression")
        
    if WORKFLOW_AVAILABLE:
            display_workflow_progress(workflow)
        
        # Bouton de téléchargement partiel (si chapitres validés)
    progress = workflow.get_progress()
    if progress['validated_chapters'] > 0:
            st.markdown("---")
            st.markdown("### [SAVE] Sauvegarde intermédiaire")
            
            st.info(f"💡 Vous avez {progress['validated_chapters']} chapitre(s) validé(s). Vous pouvez les télécharger dès maintenant.")
            
            # Compiler les chapitres validés seulement
            partial_report = ""
            for chapter in workflow.chapters:
                if chapter.status.value == "Validé" and chapter.content:
                    partial_report += chapter.content + "\n\n---\n\n"
            
            if partial_report:
                st.markdown("#### 📥 Télécharger les chapitres validés")
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    # Markdown
                    st.download_button(
                        "Markdown",
                        partial_report,
                        file_name=f"rapport_partiel_{progress['validated_chapters']}_chapitres_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                with col2:
                    # HTML
                    try:
                        import re
                        html_body = partial_report
                        html_body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
                        html_body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
                        html_body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
                        html_body = html_body.replace('\n\n', '</p><p>')
                        html_body = '<p>' + html_body + '</p>'
                        html_body = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1" />', html_body)
                        html_body = html_body.replace('<p></p>', '').replace('<p><h', '<h').replace('</h1></p>', '</h1>').replace('</h2></p>', '</h2>').replace('</h3></p>', '</h3>')
                        
                        html_content = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Rapport Partiel</title>
<style>
body {{ font-family: Arial; max-width: 900px; margin: 40px auto; padding: 20px; }}
h1, h2, h3 {{ color: #2c3e50; }}
img {{ max-width: 100%; height: auto; margin: 15px 0; }}
</style></head><body>{html_body}</body></html>"""
                        
                        st.download_button(
                            "HTML",
                            html_content,
                            file_name=f"rapport_partiel_{progress['validated_chapters']}_chapitres_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                            mime="text/html",
                            use_container_width=True
                        )
                    except:
                        st.warning("HTML indisponible")
                
                with col3:
                    # Word
                    try:
                        from docx import Document
                        import io
                        import re
                        
                        doc = Document()
                        doc.add_heading('Rapport Partiel', 0)
                        doc.add_paragraph(f"{progress['validated_chapters']} chapitres validés")
                        doc.add_paragraph()  # Ligne vide
                        
                        # Parser basique du Markdown
                        lines = partial_report.split('\n')
                        for line in lines:
                            try:
                                if line.startswith('# '):
                                    doc.add_heading(line[2:], level=1)
                                elif line.startswith('## '):
                                    doc.add_heading(line[3:], level=2)
                                elif line.startswith('### '):
                                    doc.add_heading(line[4:], level=3)
                                elif line.strip() and not line.startswith('|'):
                                    doc.add_paragraph(line)
                            except:
                                # Si erreur, ajouter comme texte simple
                                if line.strip():
                                    doc.add_paragraph(line)
                        
                        docx_bytes = io.BytesIO()
                        doc.save(docx_bytes)
                        docx_bytes.seek(0)
                        
                        st.download_button(
                            "Word",
                            docx_bytes.getvalue(),
                            file_name=f"rapport_partiel_{progress['validated_chapters']}_chapitres_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.warning(f"Word indisponible")
                
                with col4:
                    st.info("PDF au final")
        
    st.markdown("---")
        
        # Chapitre en cours
    current_chapter = workflow.get_current_chapter()
        
    if current_chapter is None:
            st.success("🎉 Tous les chapitres ont été générés et validés !")
            
            if st.button("📄 Compiler le rapport final", type="primary"):
                with st.spinner("Compilation en cours..."):
                    final_report = workflow.compile_report()
                    
                    st.success("[OK] Rapport compilé avec succès !")
                    
                    # Afficher le rapport
                    with st.expander("📄 Aperçu du rapport final", expanded=True):
                        st.markdown(final_report)
                    
                    # Options de téléchargement
                    st.markdown("### 📥 Télécharger le rapport")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown
                        st.download_button(
                            "Markdown",
                            final_report,
                            file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )
                    
                    with col2:
                        # HTML - Version améliorée avec support des tableaux
                        try:
                            import re
                            
                            def markdown_to_html(md_text):
                                """
                                Convertit Markdown en HTML de manière professionnelle 
                                avec support complet des tableaux, code et images base64.
                                """
                                import markdown
                                import re

                                # Nettoyage préliminaire : s'assurer qu'il y a des lignes vides autour des tableaux
                                # pour que le parser les détecte correctement.
                                md_text = re.sub(r'([^\n])\n\|', r'\1\n\n|', md_text)
                                md_text = re.sub(r'\|\n([^\n])', r'|\n\n\1', md_text)

                                # Utilisation du parser officiel avec extensions
                                # 'tables' : gère les tableaux avec |
                                # 'fenced_code' : gère les blocs avec ```
                                # 'nl2br' : gère les sauts de ligne naturels
                                extensions = ['tables', 'fenced_code', 'nl2br', 'attr_list']
                                
                                html_content = markdown.markdown(md_text, extensions=extensions)

                                # Post-traitement des images base64 pour s'assurer qu'elles sont responsives
                                # On remplace les balises img simples par des versions stylisées
                                html_content = html_content.replace(
                                    '<img ', 
                                    '<img style="max-width:100%; height:auto; border-radius:8px; margin:20px 0; box-shadow: 0 4px 8px rgba(0,0,0,0.1);" '
                                )

                                return html_content
                            
                            def convert_md_table_to_html(table_lines):
                                """Convertit un tableau Markdown en HTML"""
                                if len(table_lines) < 2:
                                    return '\n'.join(table_lines)
                                
                                html = '<table>\n'
                                
                                # Header (première ligne)
                                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                                html += '<thead><tr>'
                                for cell in header_cells:
                                    html += f'<th>{cell}</th>'
                                html += '</tr></thead>\n'
                                
                                # Ignorer la ligne de séparation (---) qui est la ligne 1
                                # Données (à partir de la ligne 2)
                                html += '<tbody>\n'
                                for line in table_lines[2:]:
                                    if '|' in line:
                                        cells = [c.strip() for c in line.split('|') if c.strip()]
                                        html += '<tr>'
                                        for cell in cells:
                                            html += f'<td>{cell}</td>'
                                        html += '</tr>\n'
                                html += '</tbody>\n</table>\n'
                                
                                return html
                            
                            html_body = markdown_to_html(final_report)
                            
                            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Rapport Statistique</title>
    <style>
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1, h2, h3, h4 {{ color: #2c3e50; }}
        h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; margin-top: 30px; }}
        h2 {{ border-bottom: 2px solid #95a5a6; padding-bottom: 8px; margin-top: 25px; }}
        h3 {{ margin-top: 20px; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            font-size: 14px;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        tr:hover {{ background-color: #f0f0f0; }}
        img {{
            max-width: 100%;
            height: auto;
            margin: 20px 0;
            border: 1px solid #ddd;
            border-radius: 4px;
            display: block;
        }}
        p {{ margin: 10px 0; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>
"""
                            st.download_button(
                                "HTML",
                                html_content,
                                file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                                mime="text/html",
                                use_container_width=True
                            )
                        except Exception as e:
                            st.error(f"Erreur HTML: {str(e)[:50]}")
                    
                    with col3:
                        # Word
                        try:
                            from docx import Document
                            from docx.shared import Inches, Pt, RGBColor
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            import io
                            import re
                            import base64
                            
                            doc = Document()
                            
                            # Style du document
                            style = doc.styles['Normal']
                            style.font.name = 'Calibri'
                            style.font.size = Pt(11)
                            
                            # Parser le Markdown et convertir en Word
                            lines = final_report.split('\n')
                            i = 0
                            while i < len(lines):
                                line = lines[i]
                                
                                # Titres
                                if line.startswith('# '):
                                    p = doc.add_heading(line[2:], level=1)
                                elif line.startswith('## '):
                                    p = doc.add_heading(line[3:], level=2)
                                elif line.startswith('### '):
                                    p = doc.add_heading(line[4:], level=3)
                                
                                # Images base64
                                elif line.startswith('!['):
                                    match = re.match(r'!\[([^\]]*)\]\(data:image/png;base64,([^)]+)\)', line)
                                    if match:
                                        try:
                                            img_b64 = match.group(2)
                                            img_data = base64.b64decode(img_b64)
                                            
                                            # Ajouter l'image au document
                                            img_stream = io.BytesIO(img_data)
                                            doc.add_picture(img_stream, width=Inches(6))
                                            
                                            # Ajouter légende
                                            if match.group(1):
                                                p = doc.add_paragraph(match.group(1))
                                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                p.runs[0].italic = True
                                        except:
                                            pass
                                
                                # Tableaux Markdown
                                elif '|' in line and i+1 < len(lines) and '---' in lines[i+1]:
                                    # Parser le tableau
                                    try:
                                        parts = line.split('|')
                                        if len(parts) >= 3:  # Au moins | col1 | col2 |
                                            headers = [h.strip() for h in parts[1:-1] if h.strip()]
                                            i += 2  # Skip separator
                                            
                                            if not headers:  # Sécurité
                                                i += 1
                                                continue
                                            
                                            rows = []
                                            while i < len(lines) and '|' in lines[i]:
                                                row_parts = lines[i].split('|')
                                                if len(row_parts) >= 3:
                                                    row = [c.strip() for c in row_parts[1:-1] if c.strip()]
                                                    # S'assurer que le nombre de cellules correspond
                                                    if len(row) == len(headers):
                                                        rows.append(row)
                                                i += 1
                                            
                                            # Créer le tableau Word
                                            if rows and headers:
                                                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                                                table.style = 'Light Grid Accent 1'
                                                
                                                # Headers
                                                for j, header in enumerate(headers):
                                                    if j < len(table.rows[0].cells):
                                                        table.rows[0].cells[j].text = header
                                                
                                                # Rows
                                                for r_idx, row in enumerate(rows):
                                                    if r_idx + 1 < len(table.rows):
                                                        for c_idx, cell in enumerate(row):
                                                            if c_idx < len(table.rows[r_idx+1].cells):
                                                                table.rows[r_idx+1].cells[c_idx].text = cell
                                            
                                            i -= 1
                                    except Exception as e:
                                        # Si erreur de parsing du tableau, ignorer
                                        pass
                                
                                # Texte normal
                                elif line.strip():
                                    doc.add_paragraph(line)
                                
                                i += 1
                            
                            # Sauvegarder en bytes
                            docx_bytes = io.BytesIO()
                            doc.save(docx_bytes)
                            docx_bytes.seek(0)
                            
                            st.download_button(
                                "Word",
                                docx_bytes.getvalue(),
                                file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        except ImportError:
                            st.warning("python-docx non installé")
                        except Exception as e:
                            st.error(f"Erreur Word: {str(e)[:50]}")
                    
                    with col4:
                        # PDF
                        try:
                            import re
                            
                            # UTILISER LA MÊME FONCTION QUE POUR HTML
                            def markdown_to_html_pdf(md_text):
                                """Convertit Markdown en HTML avec support des tableaux"""
                                html = md_text
                                
                                # 1. Convertir les tableaux Markdown
                                lines = html.split('\n')
                                in_table = False
                                result_lines = []
                                table_lines = []
                                
                                for line in lines:
                                    if '|' in line and not in_table:
                                        in_table = True
                                        table_lines = [line]
                                    elif '|' in line and in_table:
                                        table_lines.append(line)
                                    elif in_table:
                                        # Fin du tableau, convertir
                                        result_lines.append(convert_md_table_to_html_pdf(table_lines))
                                        result_lines.append(line)
                                        in_table = False
                                        table_lines = []
                                    else:
                                        result_lines.append(line)
                                
                                if in_table:
                                    result_lines.append(convert_md_table_to_html_pdf(table_lines))
                                
                                html = '\n'.join(result_lines)
                                
                                # 2. Convertir les headers
                                html = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
                                html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
                                html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
                                html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
                                
                                # 3. Convertir les images
                                html = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', r'<img src="\2" alt="\1" style="max-width:100%; height:auto;" />', html)
                                
                                # 4. Convertir les paragraphes
                                lines = html.split('\n')
                                result = []
                                for line in lines:
                                    if line.strip() and not line.startswith('<'):
                                        result.append(f'<p>{line}</p>')
                                    else:
                                        result.append(line)
                                html = '\n'.join(result)
                                
                                return html
                            
                            def convert_md_table_to_html_pdf(table_lines):
                                """Convertit un tableau Markdown en HTML"""
                                if len(table_lines) < 2:
                                    return '\n'.join(table_lines)
                                
                                html = '<table>\n'
                                
                                # Header
                                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                                html += '<thead><tr>'
                                for cell in header_cells:
                                    html += f'<th>{cell}</th>'
                                html += '</tr></thead>\n'
                                
                                # Données
                                html += '<tbody>\n'
                                for line in table_lines[2:]:
                                    if '|' in line:
                                        cells = [c.strip() for c in line.split('|') if c.strip()]
                                        html += '<tr>'
                                        for cell in cells:
                                            html += f'<td>{cell}</td>'
                                        html += '</tr>\n'
                                html += '</tbody>\n</table>\n'
                                
                                return html
                            
                            html_body = markdown_to_html_pdf(final_report)
                            
                            html_for_pdf = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            margin: 2cm;
        }}
        body {{ 
            font-family: Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #333;
        }}
        h1, h2, h3, h4 {{ 
            color: #2c3e50;
            page-break-after: avoid;
        }}
        h1 {{ 
            font-size: 20pt;
            border-bottom: 2px solid #3498db;
            padding-bottom: 8px;
            margin-top: 20px;
        }}
        h2 {{ 
            font-size: 16pt;
            border-bottom: 1px solid #95a5a6;
            padding-bottom: 6px;
            margin-top: 18px;
        }}
        h3 {{ 
            font-size: 14pt;
            margin-top: 15px;
        }}
        h4 {{ 
            font-size: 12pt;
            margin-top: 12px;
        }}
        table {{ 
            border-collapse: collapse;
            width: 100%;
            margin: 12px 0;
            page-break-inside: avoid;
        }}
        th, td {{ 
            border: 1px solid #999;
            padding: 6px 10px;
            text-align: left;
            font-size: 10pt;
        }}
        th {{ 
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{ 
            background-color: #f5f5f5;
        }}
        img {{ 
            max-width: 100%;
            height: auto;
            margin: 12px 0;
            page-break-inside: avoid;
        }}
        p {{ 
            margin: 8px 0;
            text-align: justify;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>
"""
                            from weasyprint import HTML as WeasyHTML
                            pdf_bytes = WeasyHTML(string=html_for_pdf).write_pdf()
                            
                            st.download_button(
                                "PDF",
                                pdf_bytes,
                                file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        except ImportError as e:
                            st.warning("weasyprint non installé (nécessite des bibliothèques système)")
                        except Exception as e:
                            st.warning(f"PDF non disponible")
        
    else:
            st.markdown(f"### 📝 Chapitre {current_chapter.number} : {current_chapter.title}")
            
            # Afficher les sections du chapitre
            with st.expander("📋 Sections à couvrir", expanded=False):
                for i, section in enumerate(current_chapter.sections, 1):
                    st.markdown(f"**{current_chapter.number}.{i}. {section['titre']}**")
                    for analyse in section['analyses']:
                        st.markdown(f"  • {analyse}")
            
            # Bouton de génération
            if current_chapter.status.value in ["En attente", "Rejeté (à regénérer)", "Erreur"]:
                
                if st.button(f"🚀 Générer le Chapitre {current_chapter.number}", type="primary"):
                    
                    # Créer un placeholder pour les étapes
                    status_placeholder = st.empty()
                    progress_placeholder = st.empty()
                    
                    # Étape 1
                    progress_placeholder.progress(0.2, "Étape 1/5: Chargement du contexte...")
                    status_placeholder.info("📚 Chargement du contexte des chapitres précédents")
                    import time
                    time.sleep(0.5)
                    
                    # Étape 2
                    progress_placeholder.progress(0.4, "Étape 2/5: Construction du prompt IA")
                    status_placeholder.info("📝 Construction du prompt IA")
                    time.sleep(0.5)
                    
                    # Étape 3
                    progress_placeholder.progress(0.6, "Étape 4/6: Génération du contenu...")
                    status_placeholder.info("🤖 Génération du contenu")
                    
                    result = workflow.generate_current_chapter()
                    
                    if result['success']:
                        # Étape 4
                        progress_placeholder.progress(0.8, "Étape 5/6: Exécution du code Python...")
                        status_placeholder.info("⚙️ Exécution du code Python dans E2B")
                        time.sleep(0.5)
                        
                        # Étape 5
                        progress_placeholder.progress(1.0, "Étape 6/6: Finalisation...")
                        status_placeholder.info("[OK] Finalisation du chapitre")
                        time.sleep(0.5)
                        
                        # Nettoyer
                        progress_placeholder.empty()
                        status_placeholder.empty()
                        
                        st.success(f"[OK] Chapitre {current_chapter.number} généré avec succès !")
                        
                        if LOGGING_AVAILABLE:
                            log_user_action('chapter_generated', {
                                'chapter_number': current_chapter.number,
                                'word_count': len(result['content'].split())
                            })
                        
                        time.sleep(1)
                        st.rerun()
                    
                    else:
                        progress_placeholder.empty()
                        status_placeholder.empty()
                        
                        st.error(f"[ERROR] Erreur lors de la génération : {result['error']}")
                        
                        if LOGGING_AVAILABLE:
                            logger.error(f"Chapter generation failed: {result['error']}")
            
            # Afficher le chapitre si généré
            if current_chapter.status.value == "Généré (en attente validation)":
                
                st.markdown("### 📄 Contenu généré")
                
                # Option : Mode édition
                with st.expander("✏️ Modifier le contenu", expanded=False):
                    st.info("💡 Vous pouvez modifier le texte ci-dessous avant de valider le chapitre")
                    
                    # Initialiser le contenu original en session
                    if f'original_content_{current_chapter.number}' not in st.session_state:
                        st.session_state[f'original_content_{current_chapter.number}'] = current_chapter.content
                    
                    # Zone d'édition
                    edited_content = st.text_area(
                        "Contenu du chapitre :",
                        value=current_chapter.content,
                        height=400,
                        key=f"edit_chapter_{current_chapter.number}"
                    )
                    
                    # Détecter si des modifications ont été faites
                    original_content = st.session_state[f'original_content_{current_chapter.number}']
                    has_changes = edited_content != original_content
                    
                    # Afficher le nombre de modifications
                    if has_changes:
                        char_diff = len(edited_content) - len(original_content)
                        if char_diff > 0:
                            st.caption(f"📝 {abs(char_diff)} caractères ajoutés")
                        elif char_diff < 0:
                            st.caption(f"📝 {abs(char_diff)} caractères supprimés")
                        else:
                            st.caption(f"📝 Contenu modifié")
                    
                    # Boutons d'action
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Bouton Enregistrer (vert si modifications, gris sinon)
                        save_disabled = not has_changes
                        
                        if st.button(
                            "[SAVE] Enregistrer les modifications",
                            key=f"save_edit_{current_chapter.number}",
                            type="primary" if has_changes else "secondary",
                            disabled=save_disabled,
                            use_container_width=True,
                            help="Enregistrer les modifications apportées" if has_changes else "Aucune modification à enregistrer"
                        ):
                            # Appliquer les modifications
                            current_chapter.content = edited_content
                            st.session_state[f'original_content_{current_chapter.number}'] = edited_content
                            
                            # Message de succès
                            st.success("[OK] Modifications enregistrées avec succès !")
                            
                            if LOGGING_AVAILABLE:
                                log_user_action('chapter_edited', {
                                    'chapter_number': current_chapter.number,
                                    'char_diff': len(edited_content) - len(original_content)
                                })
                            
                            # Attendre un peu pour que l'utilisateur voie le message
                            import time
                            time.sleep(1)
                            st.rerun()
                    
                    with col2:
                        # Bouton Annuler (rouge si modifications, gris sinon)
                        cancel_disabled = not has_changes
                        
                        if st.button(
                            "[REFRESH] Annuler les modifications",
                            key=f"cancel_edit_{current_chapter.number}",
                            disabled=cancel_disabled,
                            use_container_width=True,
                            help="Annuler et revenir au contenu original" if has_changes else "Aucune modification à annuler"
                        ):
                            # Réinitialiser au contenu original
                            current_chapter.content = original_content
                            
                            # Message d'info
                            st.warning("↩️ Modifications annulées")
                            
                            if LOGGING_AVAILABLE:
                                log_user_action('chapter_edit_cancelled', {
                                    'chapter_number': current_chapter.number
                                })
                            
                            import time
                            time.sleep(0.5)
                            st.rerun()
                
                # Aperçu du contenu (lecture seule)
                st.markdown("**Aperçu :**")
                with st.container():
                    # Afficher le contenu avec support des images base64 inline
                    import re
                    import base64
                    from io import BytesIO
                    from PIL import Image
                    
                    content = current_chapter.content
                    
                    # Pattern pour détecter les images base64 inline
                    base64_pattern = r'!\[([^\]]*)\]\(data:image/png;base64,([^)]+)\)'
                    
                    # Diviser le contenu en sections (texte et images)
                    parts = re.split(base64_pattern, content)
                    
                    # Afficher chaque partie
                    for i, part in enumerate(parts):
                        if i % 3 == 0:
                            # Texte normal
                            if part.strip():
                                st.markdown(part, unsafe_allow_html=True)
                        elif i % 3 == 1:
                            # Titre de l'image (on l'ignore ou on pourrait l'afficher)
                            pass
                        elif i % 3 == 2:
                            # Données base64 de l'image
                            try:
                                # Décoder le base64
                                img_data = base64.b64decode(part)
                                img = Image.open(BytesIO(img_data))
                                
                                # Afficher l'image avec Streamlit
                                st.image(img, use_container_width=True)
                            except Exception as e:
                                st.error(f"Erreur d'affichage de l'image: {e}")
                    
                    # Si pas d'images base64, afficher normalement
                    if not re.search(base64_pattern, content):
                        st.markdown(content, unsafe_allow_html=True)
                
                st.markdown("---")
                st.markdown("### [OK] Validation du chapitre")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button("[OK] Valider ce chapitre", type="primary", use_container_width=True, help="Valider le chapitre et passer au suivant"):
                        workflow.validate_chapter(current_chapter.number)
                        
                        st.success(f"[OK] Chapitre {current_chapter.number} validé !")
                        
                        if LOGGING_AVAILABLE:
                            log_user_action('chapter_validated', {
                                'chapter_number': current_chapter.number
                            })
                        
                        st.rerun()
                
                with col2:
                    # Nouveau système de régénération avec feedback
                    if f'show_regen_form_{current_chapter.number}' not in st.session_state:
                        st.session_state[f'show_regen_form_{current_chapter.number}'] = False
                    
                    if not st.session_state[f'show_regen_form_{current_chapter.number}']:
                        # Bouton pour ouvrir le formulaire
                        if st.button("[REFRESH] Modifier et régénérer", use_container_width=True, help="Demander des modifications spécifiques"):
                            st.session_state[f'show_regen_form_{current_chapter.number}'] = True
                            st.rerun()
                    else:
                        # Afficher le formulaire de modifications
                        st.markdown("---")
                        st.markdown("### [REFRESH] Instructions de régénération")
                        
                        st.info("""
💡 **Soyez précis dans vos demandes !**

**Exemples de modifications :**
- "Ajouter plus de graphiques (au moins 3)"
- "Rendre l'analyse plus concise (2 pages maximum)"
- "Approfondir l'analyse des corrélations entre X et Y"
- "Ajouter un test statistique pour vérifier H0"
- "Simplifier le vocabulaire pour un public non-technique"
- "Ajouter une comparaison avec les résultats de l'étude Z"
                        """)
                        
                        # Zone de texte pour les instructions
                        modification_request = st.text_area(
                            "📝 Que souhaitez-vous améliorer ou modifier ?",
                            height=150,
                            placeholder="Décrivez précisément les modifications souhaitées...",
                            help="Plus vous êtes précis, meilleure sera la régénération",
                            key=f"regen_request_{current_chapter.number}"
                        )
                        
                        # Options supplémentaires
                        col_opt1, col_opt2 = st.columns(2)
                        
                        with col_opt1:
                            keep_structure = st.checkbox(
                                "Conserver la structure actuelle",
                                value=True,
                                help="Garde les mêmes sections, modifie seulement le contenu",
                                key=f"keep_struct_{current_chapter.number}"
                            )
                        
                        with col_opt2:
                            keep_code = st.checkbox(
                                "Conserver les analyses existantes",
                                value=False,
                                help="Garde les calculs actuels, améliore seulement la rédaction",
                                key=f"keep_code_{current_chapter.number}"
                            )
                        
                        # Boutons d'action
                        col_a, col_b = st.columns(2)
                        
                        with col_a:
                            if st.button("[REFRESH] Régénérer avec ces modifications", type="primary", use_container_width=True):
                                if modification_request.strip():
                                    # Sauvegarder les instructions dans le chapitre
                                    current_chapter.regeneration_instructions = {
                                        'request': modification_request,
                                        'keep_structure': keep_structure,
                                        'keep_code': keep_code
                                    }
                                    
                                    # Régénérer avec les instructions
                                    workflow.reject_chapter(current_chapter.number, modification_request)
                                    
                                    st.success("[OK] Instructions enregistrées ! Régénération en cours...")
                                    
                                    if LOGGING_AVAILABLE:
                                        log_user_action('chapter_regeneration_requested', {
                                            'chapter_number': current_chapter.number,
                                            'request': modification_request[:100],
                                            'keep_structure': keep_structure,
                                            'keep_code': keep_code
                                        })
                                    
                                    # Fermer le formulaire
                                    st.session_state[f'show_regen_form_{current_chapter.number}'] = False
                                    st.rerun()
                                else:
                                    st.warning("[WARNING] Veuillez décrire les modifications souhaitées")
                        
                        with col_b:
                            if st.button("[ERROR] Annuler", use_container_width=True):
                                st.session_state[f'show_regen_form_{current_chapter.number}'] = False
                                st.rerun()



    # Navigation workflow
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← Retour config", use_container_width=True, key="back_gen_wf"):
            st.session_state.workflow_step = 4
            st.rerun()
    with col3:
        if st.button("[REFRESH] Nouveau rapport", use_container_width=True, key="restart_wf"):
            st.session_state.workflow_step = 1
            st.session_state.workflow_history = [1]
            st.session_state.plan = None
            st.session_state.analysis_mode = None
            st.rerun()

else:
    st.error(f"[ERROR] Étape inconnue : {current_step}")
    if st.button("[REFRESH] Recommencer", type="primary"):
        st.session_state.workflow_step = 1
        st.session_state.workflow_history = [1]
        st.rerun()